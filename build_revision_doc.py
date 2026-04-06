from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("Distributed_Systems_3_Week_Exam_Revision_Guide.docx")

PAPERS_ANALYSED = [
    "2022-2023 main paper",
    "2022-2023 reassessment paper",
    "2023-2024 main paper",
    "2023-2024 refer/defer paper",
    "2024-2025 main paper",
    "2024-2025 resit paper",
]

REPEATED_TOPICS = [
    [
        "Security: HTTPS, TLS, certificates, signatures",
        "6/6 papers",
        "Step-by-step trust, handshake stages, certificate chain, digital signatures, and why HTTPS is safer than HTTP.",
        "Very high",
    ],
    [
        "Messaging and inter-service communication",
        "6/6 papers",
        "HTTP vs queue vs pub-sub vs event-driven approaches, plus synchronous and asynchronous trade-offs.",
        "Very high",
    ],
    [
        "Architecture: monolith vs microservices and distributed design",
        "6/6 papers",
        "Why organisations split systems, when that helps, and what complexity it introduces.",
        "Very high",
    ],
    [
        "API fundamentals: endpoint design, Express, middleware, JSON, CORS, OpenAPI",
        "6/6 papers",
        "Clean resource design, request handling, response codes, browser rules and documentation.",
        "Very high",
    ],
    [
        "Cloud models, resilience, AZs, elastic compute, load balancing",
        "5/6 papers",
        "IaaS/PaaS/SaaS/DBaaS, shared responsibility, resilient architecture and scaling components.",
        "High",
    ],
    [
        "Containers, Docker and Kubernetes",
        "5/6 papers",
        "Container benefits, persistence, environment variables, Kubernetes objects and rolling updates.",
        "High",
    ],
    [
        "Identity and access: IAM, OAuth, OIDC, JWT, SSO",
        "4/6 papers",
        "Authentication vs authorisation, token structure, and login sequence questions.",
        "High",
    ],
    [
        "Data and scaling",
        "3-5/6 papers",
        "Referential integrity, connection pools, read replicas, sharding, consistent hashing and transaction limits.",
        "Medium-high",
    ],
    [
        "Terraform and Infrastructure as Code",
        "1/6 papers",
        "Low recurrence so far, but it is clearly present in the current slide deck and recent exam.",
        "Medium",
    ],
]

COMMAND_WORDS = [
    [
        "Explain",
        "This is the dominant command word by far: the examiner wants a clear definition, then a process, then why it matters.",
        "1. Define the term. 2. Walk through the flow or mechanism in order. 3. Add why it is useful or important. 4. Add a short example if possible.",
        "Writing only a definition, missing the sequence, or giving a generic paragraph that never answers 'how'.",
    ],
    [
        "Compare",
        "Show both sides against the same criteria so the differences are obvious.",
        "1. State the two things. 2. Compare them point-by-point under the same headings. 3. End with when one is preferable.",
        "Writing two separate mini-essays with no direct comparison, or listing features without saying which is better and why.",
    ],
    [
        "Discuss",
        "Give a balanced answer that considers strengths, limits and context.",
        "1. State the issue. 2. Cover both benefits and drawbacks. 3. Link them to a use-case. 4. Finish with a reasoned conclusion.",
        "Being one-sided, staying too general, or forgetting to tie the discussion back to the scenario.",
    ],
    [
        "Evaluate",
        "Judge an option against criteria such as security, cost, complexity, resilience or performance.",
        "1. State the criteria. 2. Weigh the evidence on both sides. 3. Give a justified judgement at the end.",
        "Describing a technology without judging it, or giving a verdict with no criteria.",
    ],
    [
        "Justify",
        "Choose an answer and defend it with reasons linked to the use-case.",
        "1. Make the choice. 2. Give two or three precise reasons. 3. Link each reason to the scenario or requirement.",
        "Making a choice with no reasoning, or using stock advantages that could apply to anything.",
    ],
    [
        "Design",
        "Propose a sensible endpoint, architecture or diagram and explain why the design is formed that way.",
        "1. State the design. 2. Label the important parts. 3. Explain each choice. 4. Give an example call, data flow or use-case.",
        "Giving only the final diagram or URL and not explaining the rationale behind it.",
    ],
]

REVISION_PLAN = [
    {
        "day": "Day 1",
        "topics": "Distributed systems basics, monoliths, microservices, functional vs non-functional requirements",
        "order": "Start with what a distributed system is, then monolith vs microservices, then requirement types.",
        "goal": "Understand why systems are split up and when they should not be.",
        "recall": "Write five differences between a monolith and a microservice architecture from memory.",
        "exam": "Answer a short compare question on monolith vs microservices in 20 minutes.",
    },
    {
        "day": "Day 2",
        "topics": "Microservice trade-offs and API gateway",
        "order": "Review microservice benefits first, then challenges, then API gateway role.",
        "goal": "Be able to justify a distributed design for a business use-case.",
        "recall": "List three technical and three business reasons to use microservices.",
        "exam": "Write a 15-mark style answer on why an API gateway may be recommended.",
    },
    {
        "day": "Day 3",
        "topics": "JSON, URL structure, endpoints, REST, HTTP methods",
        "order": "JSON first, then URL parts, then endpoint design, then REST method choices.",
        "goal": "Be able to design clean resource-based API calls.",
        "recall": "Write one example URL that uses both path parameters and query parameters correctly.",
        "exam": "Design an endpoint for a realistic use-case and explain your choices.",
    },
    {
        "day": "Day 4",
        "topics": "Express request handling, routes, middleware, response codes",
        "order": "Request flow first, then routes, then middleware stack, then common status codes.",
        "goal": "Follow a browser request all the way to the correct endpoint and response.",
        "recall": "Explain the middleware stack in four sentences with no notes.",
        "exam": "Do a 2024-2025 style question on how Express handles a request and what codes may return.",
    },
    {
        "day": "Day 5",
        "topics": "CORS, OpenAPI, design-first/API-first/code-first, HATEOAS light review",
        "order": "Origins and CORS first, then API documentation approaches, then HATEOAS as a secondary extra.",
        "goal": "Understand browser restrictions and API contract documentation.",
        "recall": "Write the formula for an origin and five top-level OpenAPI elements.",
        "exam": "Compare design-first, API-first and code-first for one organisation and one small team.",
    },
    {
        "day": "Day 6",
        "topics": "Relational basics, referential integrity, connection pooling",
        "order": "Primary/foreign keys first, then integrity rules, then connection pooling.",
        "goal": "Understand how the data layer stays consistent and supports concurrent users.",
        "recall": "Define cascade, restrict and connection pool in one line each.",
        "exam": "Write a short answer on why connection pooling is better than one direct connection per request.",
    },
    {
        "day": "Day 7",
        "topics": "Cloud business model, IaaS/PaaS/SaaS/DBaaS, shared responsibility",
        "order": "On-prem vs cloud first, then service models, then responsibility split.",
        "goal": "Choose the right cloud model for a scenario and defend the choice.",
        "recall": "Draw a quick ladder from on-prem to SaaS and note who manages more at each step.",
        "exam": "Answer an IaaS vs PaaS question with a use-case and shared responsibility reference.",
    },
    {
        "day": "Day 8",
        "topics": "Resilience, high availability, availability zones, load balancers, elastic compute",
        "order": "Define resilience and HA first, then AZs, then load balancers, then auto-scaling and elastic compute.",
        "goal": "Be able to explain a resilient cloud design with labelled components.",
        "recall": "Explain the difference between resilience and high availability without using the words 'same thing'.",
        "exam": "Answer a diagram question on a load balancer in an elastic compute solution.",
    },
    {
        "day": "Day 9",
        "topics": "VMs, hypervisors, containers, Docker images and containers",
        "order": "VM basics first, then hypervisors, then containers, then image vs container.",
        "goal": "Understand the deployment model before moving to Docker Compose and Kubernetes.",
        "recall": "List four differences between a VM and a container.",
        "exam": "Do a compare question on VM vs container with one use-case for each.",
    },
    {
        "day": "Day 10",
        "topics": "Docker Compose, named volumes, bind mounts, environment variables",
        "order": "Compose overview first, then persistence, then configuration through env vars.",
        "goal": "Be able to explain persistence and configuration in multi-container systems.",
        "recall": "State the problem a named volume solves and what happens if you do not use one.",
        "exam": "Answer a named volume question and a short environment variable question.",
    },
    {
        "day": "Day 11",
        "topics": "Synchronous vs asynchronous communication, HTTP vs message queue",
        "order": "Direct request-response first, then queue-based messaging, then loose coupling trade-offs.",
        "goal": "Choose the right communication style for the right problem.",
        "recall": "Write three pros and three cons of synchronous HTTP between microservices.",
        "exam": "Answer a question on using a message queue to increase resilience.",
    },
    {
        "day": "Day 12",
        "topics": "Publisher-subscriber, event-driven architecture, thick vs thin events, RabbitMQ vs Kafka",
        "order": "Pub-sub first, then events, then thick vs thin, then broker comparison.",
        "goal": "Explain event flow and eventual consistency clearly.",
        "recall": "Define thick and thin events and say when you would choose each.",
        "exam": "Redraw a message-based or synchronous flow as an event-driven one and explain it.",
    },
    {
        "day": "Day 13",
        "topics": "Security basics, digital signatures, certificates, chain of trust",
        "order": "Hashing and HMAC first, then encryption, then signatures, then certificates.",
        "goal": "Separate the different crypto tools and what each is for.",
        "recall": "Write one sentence each for hash, HMAC, symmetric encryption, asymmetric encryption, digital signature and certificate.",
        "exam": "Explain how a digital signature provides authenticity, integrity and non-repudiation.",
    },
    {
        "day": "Day 14",
        "topics": "HTTPS, TLS handshake, session keys",
        "order": "HTTPS vs HTTP first, then browser trust, then handshake stages, then session keys.",
        "goal": "Give a clean step-by-step exam explanation of secure session setup.",
        "recall": "Write the TLS handshake stages from memory in numbered order.",
        "exam": "Answer a full TLS/HTTPS question using a step-by-step structure.",
    },
    {
        "day": "Day 15",
        "topics": "Authentication, authorisation, IAM, JWT",
        "order": "Auth vs authorisation first, then IAM terms, then JWT structure and claims.",
        "goal": "Know the identity vocabulary cold.",
        "recall": "Write JWT = header.payload.signature and explain what each part is.",
        "exam": "Answer a short IAM/JWT question set from an older paper.",
    },
    {
        "day": "Day 16",
        "topics": "OAuth2 and OpenID Connect authorisation code flow",
        "order": "OAuth purpose first, OIDC extension second, then the full browser/app/IdP sequence.",
        "goal": "Explain the login flow in order without mixing up the roles.",
        "recall": "Sketch the OIDC flow with browser, app server and identity provider only.",
        "exam": "Write a sequence-diagram style answer for protected resource access via OIDC.",
    },
    {
        "day": "Day 17",
        "topics": "Kubernetes fundamentals",
        "order": "Cluster and node first, then pod, then ReplicaSet, Deployment, Service, ConfigMap and Secret.",
        "goal": "Understand the Kubernetes hierarchy before rolling updates.",
        "recall": "Define cluster, node, pod, ReplicaSet, Deployment and Service in one line each.",
        "exam": "Answer a components question on Kubernetes objects and service types.",
    },
    {
        "day": "Day 18",
        "topics": "Kubernetes rolling updates, service types, Terraform, ACID and SAGA light pass",
        "order": "Rolling updates first, then NodePort/ClusterIP/LoadBalancer, then Terraform, then distributed transaction extras.",
        "goal": "Cover advanced but smaller-yield material before full practice.",
        "recall": "State the difference between declarative and imperative in one sentence.",
        "exam": "Write a short answer on rolling updates or Terraform provisioners.",
    },
    {
        "day": "Day 19",
        "topics": "Database scaling: read replicas, sharding, consistent hashing, NoSQL angle",
        "order": "Relational scaling difficulty first, then read replicas, then sharding and consistent hashing.",
        "goal": "Be able to explain partial solutions and their trade-offs.",
        "recall": "Write two reasons relational databases are hard to horizontally scale.",
        "exam": "Answer a read replica or sharding question with one benefit and one limitation.",
    },
    {
        "day": "Day 20",
        "topics": "Timed practice set A",
        "order": "Start with one API/web question, then one cloud/containers question, then one messaging question.",
        "goal": "Practise planning and writing under time pressure.",
        "recall": "Do a 15-minute brain dump of your top 20 terms before the timed set.",
        "exam": "Write three answers under strict timing and mark where you ran short on detail.",
    },
    {
        "day": "Day 21",
        "topics": "Timed practice set B and final rapid review",
        "order": "Do one security question, one identity question and one mixed architecture/data question, then final recap.",
        "goal": "Simulate exam conditions and patch remaining weak spots.",
        "recall": "Reproduce five key diagrams from memory: load balancer, API gateway, TLS, OIDC, Kubernetes rolling update.",
        "exam": "Do a final 90-minute mock or three 25-mark answers, then update your weak-spot tracker honestly.",
    },
]

TOPICS = [
    {
        "title": "Distributed Systems Foundations",
        "what": [
            "This topic explains what a distributed system is and why organisations move away from one large monolithic application towards smaller cooperating services."
        ],
        "why": [
            "It is the foundation for the whole module. If you do not understand why the system is split up, later topics such as messaging, containers, API gateways and resilience feel like disconnected facts."
        ],
        "key_ideas": [
            "A distributed system looks like one system to the user, but it is built from multiple networked parts.",
            "A monolith is usually one main deployable application, often backed by one main database.",
            "A microservice is built around a business capability and should have a clear boundary.",
            "Microservices favour loose coupling, independent deployment and independent scaling.",
            "The price of this flexibility is extra latency, harder tracing, more operational complexity and data consistency challenges.",
            "Functional requirements say what the system must do; non-functional requirements say how well it must do it.",
            "A composite microservice is a higher-level service that coordinates smaller actions to deliver one wider business feature.",
        ],
        "simple_explanation": [
            "Imagine a tiny cafe where one person takes orders, makes drinks, heats food and takes payment. That is like a monolith: one place to change, one place to debug, one big unit to deploy.",
            "Now imagine a large airport food court. Separate stations handle payment, drinks, hot food and collection. That is closer to a distributed microservice system. If the drinks queue becomes too long, you can add another drinks station without rebuilding the whole food court, but the stations must communicate properly.",
        ],
        "example": [
            "An online retailer may split into services such as catalogue, orders, payments, inventory and shipping. The user still sees one shopping site, but behind the scenes several services cooperate to complete the purchase."
        ],
        "common_mistakes": [
            "Saying microservices are always better than monoliths.",
            "Ignoring business reasons and writing only technical points.",
            "Forgetting that network calls are slower and less reliable than in-process function calls.",
            "Assuming all services can share one database forever without design problems.",
        ],
        "how_in_exams": [
            "2022-2023 reassessment asked for a comparison of monolithic and microservice-based systems from both technical and business viewpoints.",
            "2023-2024 asked for the rationale and key properties of a microservice and a composite microservice.",
            "2024-2025 resit asked for business benefits of a distributed microservice architecture over a monolith.",
            "The papers repeatedly reward balanced trade-offs rather than one-sided praise.",
        ],
        "examiner": [
            "A definition of both architectures.",
            "Direct comparison using the same criteria: agility, scaling, cost, data consistency, deployment speed, fault isolation.",
            "A business context, not just textbook terms.",
            "A clear conclusion about when each approach is sensible.",
        ],
        "memory_tricks": [
            "Monolith = one deployable box.",
            "Microservices = many focused boxes.",
            "Functional = what it does; non-functional = how well it does it.",
        ],
        "quiz": [
            "What makes a distributed system still feel like one system to the user?",
            "What is the biggest operational advantage of a monolith for a small stable system?",
            "Why can microservices scale more cheaply in some cases?",
            "What does 'loose coupling' mean in simple terms?",
            "What is one risk of splitting a system into many services?",
        ],
        "quiz_answers": [
            "Because the separate parts work together behind the scenes and present one overall service or user experience.",
            "It is usually simpler to build, deploy, test and debug because everything is in one main application.",
            "Because you can often scale only the busy service instead of scaling the whole application.",
            "It means one service does not depend heavily on another service's internal details, so changes are less likely to break everything else.",
            "One risk is extra complexity, including more network communication, harder debugging and data consistency problems.",
        ],
        "exam_question": "Compare a monolithic architecture with a microservice-based distributed system for a growing online retailer. Include both technical and business points.",
        "model_answer": [
            "A monolithic architecture packages most or all of the business logic into one main application, often with one main database. Its strengths are simpler deployment, simpler debugging and strong in-process performance, so it is often a good choice for smaller systems or systems that do not change often.",
            "A microservice-based distributed system splits the retailer into business services such as catalogue, orders, payments and shipping. This improves agility because teams can change and deploy one service without redeploying everything. It also improves targeted scaling because only the busy service may need extra instances.",
            "However, microservices add network latency, monitoring difficulty, extra infrastructure and data consistency challenges. For a growing retailer that changes often, microservices are usually the better strategic choice because they support faster delivery and independent scaling, but they are only worth it if the organisation can manage the extra operational complexity.",
        ],
        "recap": [
            "Start by deciding whether the business actually needs agility and independent scaling. If not, the monolith may still be the smarter answer."
        ],
    },
    {
        "title": "API Basics: JSON, URLs, Endpoints and REST",
        "what": [
            "This topic covers the contract between systems when they communicate over the web: JSON data, URL structure, endpoint design and common HTTP methods."
        ],
        "why": [
            "A large part of the module depends on services talking to each other and to front ends. If you can design and explain a clean API, you pick up reliable marks."
        ],
        "key_ideas": [
            "An API is a contract that lets one program use another program's functionality.",
            "A URL is typically scheme://host:port/path?query.",
            "Path parameters usually identify the resource; query parameters usually filter, sort or limit results.",
            "JSON is a lightweight text format that represents structured data using key-value pairs and arrays.",
            "Common HTTP methods are GET, POST, PUT, PATCH and DELETE.",
            "REST-style APIs are resource-focused and predictable.",
            "HATEOAS is a stricter historical REST idea where responses include links to possible next actions.",
        ],
        "simple_explanation": [
            "Think of an API like ordering from a cafe till. You do not walk into the kitchen and use the cooker yourself. Instead, you use a clear interface that says what you want. The kitchen then does the work and sends a result back.",
            "The URL tells the server which resource you want, and the HTTP method tells it what kind of action you are requesting. JSON is simply the shape used to send the data in a way lots of systems understand.",
        ],
        "example": [
            "A sensible bookstore endpoint could be GET /books/hardback/dystopian?author=George%20Orwell&title=1984&maxPrice=10.99. The path identifies the main resource context and the query adds optional filters."
        ],
        "common_mistakes": [
            "Putting optional filters into the path instead of the query string.",
            "Using verbs in the endpoint when the HTTP method already expresses the action.",
            "Forgetting to give a concrete example call.",
            "Confusing JSON structure with how the endpoint itself is designed.",
        ],
        "how_in_exams": [
            "2024-2025 main asked for a bookstore endpoint design and an example call.",
            "2023-2024 included JSON and REST-related ideas through API and JWT questions.",
            "2023-2024 refer/defer asked for an API URL to return bank transactions.",
            "2022-2023 reassessment included HATEOAS as an older but still examinable API concept.",
        ],
        "examiner": [
            "Correct HTTP method choice.",
            "A clean and readable endpoint shape.",
            "A sensible split between path parameters and query parameters.",
            "A valid example request and, where relevant, a plausible JSON response shape.",
        ],
        "memory_tricks": [
            "Path points to the thing; query qualifies the thing.",
            "GET gets, POST creates, PUT replaces, PATCH partially updates, DELETE removes.",
        ],
        "quiz": [
            "What is the difference between a path parameter and a query parameter?",
            "Why is JSON so common in web APIs?",
            "Which HTTP method is most suitable for retrieving data only?",
            "When might you use PATCH instead of PUT?",
            "What is the main idea behind HATEOAS?",
        ],
        "quiz_answers": [
            "A path parameter usually identifies the resource itself, while a query parameter usually filters, sorts or limits the result.",
            "Because it is lightweight, text-based, widely supported and easy for many systems, especially JavaScript systems, to read and write.",
            "GET.",
            "When you want to update only part of a resource rather than replace the whole thing.",
            "The response should include links or actions that guide the client to what it can do next.",
        ],
        "exam_question": "Design an API endpoint to retrieve books by media type and genre, with optional filters for author, title and maximum price. Explain why your design is sensible and give an example call.",
        "model_answer": [
            "A sensible design is GET /books/{media}/{genre}?author={author}&title={title}&maxPrice={price}. I would use GET because the client is retrieving data rather than changing it. The path is used for the core resource context, while the query parameters are optional filters that may or may not be present.",
            "This design is readable, easy to extend and follows a REST-style resource approach. It also makes it clear which parts are required and which are optional. An example call is https://bookshop.example.com:443/books/hardback/dystopian?author=George%20Orwell&title=1984&maxPrice=10.99.",
            "A good response format is JSON because it is lightweight and common in web APIs. In the exam, it is worth stating both the endpoint and the reason it is structured that way, because design plus justification usually attracts the marks.",
        ],
        "recap": [
            "If the value identifies the main resource, think path. If it narrows or filters results, think query."
        ],
    },
    {
        "title": "Express Request Handling, Middleware, Response Codes and CORS",
        "what": [
            "This topic covers how a Node.js Express server receives a request, passes it through middleware, selects the right route and returns a response, including how cross-origin browser requests are controlled."
        ],
        "why": [
            "These are heavily exam-tested practical ideas. They are easy marks if you can explain the request flow in order."
        ],
        "key_ideas": [
            "Express maps URL patterns and methods to route handlers.",
            "Middleware runs between receiving the request and sending the response.",
            "Middleware is executed in the order it is registered.",
            "A middleware function can end the response or pass control onward with next().",
            "HTTP response codes summarise the result of the request.",
            "An origin is scheme + domain + port.",
            "Browsers block cross-origin requests by default unless the server's CORS policy allows them.",
        ],
        "simple_explanation": [
            "When a request arrives, Express first parses useful request information such as the method and path. It then checks middleware and routes in order until it finds the right handler.",
            "Middleware is like a checkpoint before the main destination. One checkpoint might log the request, another might check a token, and another might parse JSON. If a checkpoint decides the request should stop, it can send the response there and then.",
        ],
        "example": [
            "A PUT /users/42 request may go through logging middleware, then authentication middleware, then validation middleware, and only then reach the route handler that updates the user. If authentication fails, Express may return 401 before the update code ever runs."
        ],
        "common_mistakes": [
            "Forgetting that middleware order matters.",
            "Saying middleware only works for security.",
            "Giving response code numbers with no meaning.",
            "Confusing CORS with server-to-server authorisation. CORS is mainly a browser rule.",
        ],
        "how_in_exams": [
            "2023-2024 asked for the concept and operation of Express middleware.",
            "2024-2025 main asked how Express handles the request so the right endpoint is executed and what response codes mean.",
            "2024-2025 resit asked for the purpose and benefits of Express middleware and a use-case.",
            "2022-2023 main and 2023-2024 refer/defer both tested CORS and same-origin policy.",
        ],
        "examiner": [
            "A clear sequence from request arrival to route execution.",
            "Correct use of the terms route, middleware, req, res and next().",
            "Three sensible response codes with correct circumstances.",
            "A clear definition of origin and why CORS exists.",
        ],
        "memory_tricks": [
            "Middleware = code before the endpoint.",
            "Order matters.",
            "Origin = scheme + domain + port.",
        ],
        "quiz": [
            "What does next() do in Express middleware?",
            "Why can middleware be useful even when it does not finish the request?",
            "Give one situation where a server should return 404.",
            "What is an origin in web security terms?",
            "Why does the browser block some cross-origin requests by default?",
        ],
        "quiz_answers": [
            "It passes control to the next middleware function or route handler in the chain.",
            "Because it can still do useful work such as logging, authentication, validation or modifying the request before the endpoint runs.",
            "When the requested route or resource does not exist on the server.",
            "The scheme, domain and port taken together.",
            "To stop potentially malicious web pages reading or sending data across origins without the server's permission.",
        ],
        "exam_question": "Explain the purpose of Express middleware, how it is called during request handling, and how CORS affects a browser calling an API on another origin.",
        "model_answer": [
            "Express middleware is code that runs between receiving the HTTP request and sending the HTTP response. It can inspect or modify the request and response objects, end the request itself, or pass control to the next stage by calling next().",
            "Middleware is checked in the order it is registered. For example, a request may first be logged, then checked for a token, then validated, and finally passed to the route handler that performs the main task. This helps separate concerns because common logic does not need to be repeated inside every endpoint.",
            "CORS is a browser security mechanism. If a web page from one origin tries to call an API on another origin, the browser blocks it unless the server explicitly allows the origin, method or headers. So a good exam answer should explain both the request flow inside Express and the browser rule outside Express.",
        ],
        "recap": [
            "If the question mentions Express, think flow and order. If it mentions CORS, think browser, origin and policy."
        ],
    },
    {
        "title": "API Documentation, OpenAPI and Discoverability",
        "what": [
            "This topic explains why API documentation matters, how OpenAPI structures a contract, and when design-first, API-first or code-first documentation approaches make sense."
        ],
        "why": [
            "Past papers repeatedly test documentation because a good API is useless if nobody understands how to call it correctly."
        ],
        "key_ideas": [
            "API documentation communicates the API contract clearly and consistently.",
            "OpenAPI is the standard structure commonly used to document web APIs.",
            "Typical OpenAPI elements include info, servers, paths, components and tags.",
            "Design-first starts by modelling the API before implementation.",
            "API-first starts by formally writing the contract before coding against it.",
            "Code-first generates or builds documentation from the implemented code.",
            "HATEOAS and Richardson maturity model are lower-frequency historical extras but still worth light revision.",
        ],
        "simple_explanation": [
            "Think of API documentation as the user manual plus the exact wiring diagram for the interface. Without it, another team may guess the method, path, parameters or response shape and get it wrong.",
            "OpenAPI makes that manual structured and machine-readable. It can also power interactive documentation, mock servers and generated client code.",
        ],
        "example": [
            "In an API-first approach, a team may write an OpenAPI YAML file for /orders before backend code exists. Frontend, backend and test teams can then work in parallel using the same contract."
        ],
        "common_mistakes": [
            "Treating documentation as optional after the code is written.",
            "Listing the three approaches without saying when you would choose each one.",
            "Forgetting to mention key OpenAPI sections.",
            "Giving a very abstract answer with no organisation or project context.",
        ],
        "how_in_exams": [
            "2022-2023 reassessment asked about HATEOAS and three OpenAPI documentation approaches.",
            "2023-2024 refer/defer asked for the differences between API-first, design-first and code-first.",
            "2024-2025 resit asked why documenting API endpoints matters and what OpenAPI documentation usually contains.",
        ],
        "examiner": [
            "A clear reason why documentation matters in real teams.",
            "A correct explanation of design-first, API-first and code-first.",
            "Practical use-cases showing when each approach fits.",
            "Key OpenAPI structure elements, not just the name 'Swagger'.",
        ],
        "memory_tricks": [
            "Design-first draws first.",
            "API-first specifies first.",
            "Code-first documents the implementation you already have.",
            "OpenAPI core: Info, Servers, Paths, Components, Tags.",
        ],
        "quiz": [
            "Why can poor API documentation slow down a team even if the API works technically?",
            "What problem does OpenAPI solve?",
            "When is design-first especially useful?",
            "What is one advantage of code-first for a small team?",
            "What is the basic idea of HATEOAS?",
        ],
        "quiz_answers": [
            "Because people may guess paths, parameters or response formats, which causes mistakes, rework and wasted time.",
            "It gives a standard, clear and machine-readable way to describe the API contract.",
            "When many people or teams need to agree the interface before implementation begins.",
            "The documentation stays close to the implementation, so it can be easier to keep in sync in a smaller project.",
            "That the API response can tell the client what valid next actions are available through links or action hints.",
        ],
        "exam_question": "Compare design-first, API-first and code-first API documentation approaches and explain where you would choose each one.",
        "model_answer": [
            "Design-first begins with modelling the API before the implementation exists. It is useful in larger organisations where business and engineering teams must agree the interface early and where a shared design can reduce rework later.",
            "API-first also defines the contract before coding, but usually as a formal OpenAPI specification written in YAML or JSON. This is strong for parallel work because frontend, backend and test teams can all build against the same precise contract.",
            "Code-first starts from working code and generates or embeds documentation from it. This can be practical for smaller teams or internal systems because the documentation stays close to the implementation. The best exam answers do not just define the approaches; they explain why one would be preferred in a given project context.",
        ],
        "recap": [
            "The best documentation answer links the approach to team size, project maturity and coordination needs."
        ],
    },
    {
        "title": "Cloud Fundamentals: Service Models, Shared Responsibility and Business Value",
        "what": [
            "This topic covers what 'the cloud' means in practice, the standard cloud service models and how responsibility changes from on-premises through to managed services."
        ],
        "why": [
            "Cloud questions are common because distributed systems are usually discussed in the context of modern hosted infrastructure."
        ],
        "key_ideas": [
            "Cloud hosting replaces large up-front capital expenditure with rented operational services.",
            "IaaS gives you infrastructure to manage largely yourself.",
            "PaaS gives you a managed platform so you can focus more on the application.",
            "SaaS gives you the finished software as a service.",
            "DBaaS is a managed database service.",
            "Shared responsibility means the split of duties changes depending on the service model.",
            "Business benefits include speed, flexibility, lower up-front cost and easier scaling.",
        ],
        "simple_explanation": [
            "On-premises means you buy, house, power, secure and maintain your own infrastructure. In the cloud you rent resources from a provider that already has the buildings, hardware, networking and support model in place.",
            "The more managed the service becomes, the less control you keep but the less day-to-day responsibility you carry. That trade-off is the heart of service-model questions.",
        ],
        "example": [
            "If a team wants maximum control over the operating system, networking and custom tooling, IaaS may fit. If the team only wants to deploy code and let the platform handle the operating system and scaling, PaaS is often more appropriate."
        ],
        "common_mistakes": [
            "Saying the cloud removes all customer responsibility.",
            "Mixing up IaaS and PaaS.",
            "Giving only technical points when the question asks for business benefits.",
            "Forgetting to link the chosen model to a realistic use-case.",
        ],
        "how_in_exams": [
            "2024-2025 resit asked for IaaS vs PaaS and cloud business benefits.",
            "2022-2023 main asked for shared responsibility across service models.",
            "2023-2024 asked for cloud hosting benefits plus IaaS and PaaS properties.",
            "2022-2023 reassessment asked for IaaS, PaaS, SaaS and DBaaS with use-cases.",
        ],
        "examiner": [
            "A clear definition of each service model.",
            "A realistic use-case for the chosen model.",
            "Reference to who manages what.",
            "Business-aware reasoning, not just technical labels.",
        ],
        "memory_tricks": [
            "IaaS = I manage most of it.",
            "PaaS = Provider manages the platform.",
            "SaaS = Simply use the software.",
            "DBaaS = managed database service.",
        ],
        "quiz": [
            "What is the main difference between IaaS and PaaS?",
            "Why can the cloud reduce CapEx?",
            "What does shared responsibility mean?",
            "Give one good use-case for DBaaS.",
            "Why might a business choose PaaS over IaaS?",
        ],
        "quiz_answers": [
            "With IaaS the customer manages more of the system, while with PaaS the provider manages more of the platform for you.",
            "Because the business rents infrastructure instead of buying and maintaining large amounts of hardware up front.",
            "It means the cloud provider and the customer split the duties for security, maintenance and management.",
            "A web application that needs a database but the team does not want to run and patch database servers itself.",
            "Because PaaS can reduce operational work and let the team focus more on building the application.",
        ],
        "exam_question": "Explain the difference between IaaS and PaaS and say when you would choose one over the other with reference to shared responsibility.",
        "model_answer": [
            "IaaS provides core infrastructure such as virtual machines, storage and networking. The customer usually manages the operating system, patches, runtime, application and much of the configuration. It offers high control but also high responsibility.",
            "PaaS sits a level higher. The provider manages more of the platform, such as the operating system and much of the runtime environment, so the customer mainly focuses on deploying and configuring the application. This reduces operational effort but also reduces low-level control.",
            "I would choose IaaS when I need detailed control over the environment, for example a custom enterprise setup with specific network and operating system needs. I would choose PaaS when I want faster delivery and simpler operations, for example deploying a web API where the provider can handle scaling and patching. Shared responsibility matters because the answer is not only about features; it is about who carries which duties.",
        ],
        "recap": [
            "The more managed the service, the less you manage yourself. Always link that to a concrete use-case."
        ],
    },
    {
        "title": "Resilience, Availability Zones, Load Balancers and Elastic Compute",
        "what": [
            "This topic explains how cloud-based distributed systems stay available, recover from failure and scale with changing demand."
        ],
        "why": [
            "These ideas appear often because they connect architecture to business continuity, performance and cost."
        ],
        "key_ideas": [
            "High availability means staying up through expected failures.",
            "Resilience is wider than high availability and includes recovery and protection against larger failures.",
            "Availability zones are separate cloud locations within a region that improve fault tolerance.",
            "A load balancer spreads incoming traffic across multiple instances.",
            "Elastic compute means resources can scale up or down with demand.",
            "Horizontal scaling adds more instances; vertical scaling makes one instance larger.",
            "Auto-scaling policies often use metrics such as CPU, memory or request volume.",
        ],
        "simple_explanation": [
            "If one web server is doing all the work, it becomes a single point of failure. If it fails, the service fails. The simplest resilience improvement is to run multiple instances and place something in front of them that distributes traffic.",
            "Cloud platforms make this more powerful by letting you place instances in different availability zones and create or remove them automatically when demand changes. That is why cloud hosting is strongly linked with elastic compute.",
        ],
        "example": [
            "A web application may place a load balancer in front of several web servers spread across two availability zones. If one server fails, traffic goes to the others. If demand spikes, auto-scaling adds more instances behind the load balancer."
        ],
        "common_mistakes": [
            "Using 'resilience' and 'high availability' as if they mean exactly the same thing.",
            "Describing a load balancer as if it stores the data or does the application's business logic.",
            "Ignoring the reason for using multiple AZs.",
            "Forgetting to link scaling to actual demand changes and cost control.",
        ],
        "how_in_exams": [
            "2024-2025 resit tested resilience and load balancers directly.",
            "2023-2024 refer/defer asked about availability zones and elastic compute.",
            "2022-2023 main asked for a highly available and elastically scalable cloud architecture.",
            "Cloud architecture questions often reward diagrams very heavily.",
        ],
        "examiner": [
            "Correct definitions of resilience, high availability and elastic compute.",
            "A labelled architecture showing the role of the load balancer clearly.",
            "A direct explanation of what happens when one instance or one AZ fails.",
            "One benefit and one practical trade-off where relevant.",
        ],
        "memory_tricks": [
            "HA = stay up.",
            "Resilience = withstand and recover.",
            "Load balancer spreads traffic.",
            "Elastic compute adapts to demand.",
        ],
        "quiz": [
            "What is the difference between high availability and resilience?",
            "Why can a load balancer improve both performance and resilience?",
            "What does an availability zone protect you against?",
            "What is horizontal scaling?",
            "Why is elastic compute valuable to a business as well as technically useful?",
        ],
        "quiz_answers": [
            "High availability focuses on staying up during expected failures, while resilience is broader and includes recovery from wider disruption.",
            "Because it spreads traffic across healthy instances and can stop sending requests to failed or overloaded ones.",
            "It helps protect against the failure of a single site or data-centre location within a region.",
            "Adding more instances or machines instead of making one existing machine bigger.",
            "Because it helps the system meet demand while avoiding paying for too much idle capacity when demand drops.",
        ],
        "exam_question": "With the aid of a simple diagram, explain the purpose and operation of a load balancer in an elastic compute solution.",
        "model_answer": [
            "A load balancer is a front component that receives incoming traffic on one public address and distributes requests across multiple backend instances. In an elastic compute design, those backend instances can be virtual machines or containers that scale up and down based on demand.",
            "Its main purpose is to prevent one instance from becoming overloaded and to improve resilience. If one backend instance fails, the load balancer can stop sending traffic to it and continue using the healthy instances. If demand rises, auto-scaling can create more instances and the load balancer can start using them.",
            "In a strong exam answer, the diagram should show client traffic entering the load balancer and being spread to multiple backend nodes, ideally across more than one availability zone. The explanation should then describe both normal traffic flow and failure handling.",
        ],
        "recap": [
            "For diagram questions here, always show one front door, many backend instances, and preferably more than one zone."
        ],
    },
    {
        "title": "Containers, Docker, Compose, Volumes and Environment Variables",
        "what": [
            "This topic covers containerisation, how Docker packages applications, how Docker Compose defines multi-container solutions, and how persistence and configuration are handled."
        ],
        "why": [
            "Containers are a core deployment theme in the slide deck and the papers repeatedly test practical points such as named volumes and environment variables."
        ],
        "key_ideas": [
            "A container is an isolated runtime environment that shares the host operating system kernel.",
            "A Docker image is the packaged blueprint; a container is a running instance of that image.",
            "A virtual machine includes a full operating system; a container is lighter and starts faster.",
            "Type 1 hypervisors run directly on hardware; type 2 hypervisors run on top of a host operating system.",
            "Docker Compose defines and starts multi-container systems declaratively.",
            "Named volumes persist data outside the ephemeral container.",
            "Environment variables keep configuration out of source code and allow environment-specific settings.",
        ],
        "simple_explanation": [
            "A container is like shipping an application in its own sealed box with its dependencies. If the host machine can run containers, the box can usually be moved from one machine to another with fewer surprises than a traditional manual installation.",
            "The important limitation is that container data inside the container is usually ephemeral. If you delete the container, that writable layer disappears. That is why databases often need a named volume or another persistent storage mechanism.",
        ],
        "example": [
            "A Node API container and a MySQL container can be defined together in Docker Compose. The API reads database settings from environment variables, and the MySQL container stores its data in a named volume so the database survives container recreation."
        ],
        "common_mistakes": [
            "Saying a container is a full virtual machine.",
            "Mixing up image and container.",
            "Forgetting to explain what problem a named volume actually solves.",
            "Treating environment variables as only a security feature rather than a flexibility feature too.",
        ],
        "how_in_exams": [
            "2024-2025 main asked for benefits of containers over direct host installation.",
            "2024-2025 resit asked directly about Docker named volumes.",
            "2023-2024 asked about environment variables in containerised environments.",
            "2023-2024 refer/defer included hypervisors, containerisation and Docker Compose declarations.",
            "2022-2023 main asked for VM vs container differences.",
        ],
        "examiner": [
            "A clear statement of what problem the technology solves.",
            "One or two realistic benefits rather than a vague list.",
            "Recognition of limits, such as persistence or shared-kernel constraints.",
            "A use-case that fits the feature being discussed.",
        ],
        "memory_tricks": [
            "Image = blueprint. Container = running instance.",
            "Named volume persists. Bind mount mirrors a host folder.",
            "Env vars configure without hard-coding.",
        ],
        "quiz": [
            "What is the difference between a Docker image and a Docker container?",
            "Why can containers start faster than virtual machines?",
            "What problem does a named volume solve for a database?",
            "What is one benefit of environment variables?",
            "What is the difference between a type 1 and a type 2 hypervisor?",
        ],
        "quiz_answers": [
            "An image is the packaged blueprint; a container is a running instance created from that image.",
            "Because they do not need to boot a full guest operating system in the same way a virtual machine does.",
            "It keeps the database files outside the ephemeral container so the data survives container recreation.",
            "They let you change configuration such as ports or credentials without hard-coding values into the source code.",
            "A type 1 hypervisor runs directly on hardware, while a type 2 hypervisor runs on top of a host operating system.",
        ],
        "exam_question": "With the aid of an appropriate use-case, explain why you might use a Docker named volume and what problem it solves.",
        "model_answer": [
            "A Docker named volume is used to persist data outside an ephemeral container. This matters when the application stores important state, such as database files. If those files are kept only inside the container's writable layer, deleting or recreating the container will lose the data.",
            "A good use-case is a MySQL container in a distributed application. The container can be stopped, rebuilt or replaced, but the named volume still holds the database files on the host in Docker-managed storage. When the new container starts and mounts the same volume, the database contents remain available.",
            "Without a named volume, a database container may appear to work at first but lose all state when the container is removed. In the exam, state the problem, the solution and the consequence of not using it.",
        ],
        "recap": [
            "If the question mentions data persistence in containers, your mind should jump straight to volumes."
        ],
    },
    {
        "title": "Kubernetes Fundamentals and Rolling Updates",
        "what": [
            "This topic explains how Kubernetes orchestrates containers at scale using objects such as pods, ReplicaSets, Deployments and Services."
        ],
        "why": [
            "Kubernetes is a major current-syllabus topic and the papers test both the vocabulary and the update process."
        ],
        "key_ideas": [
            "A cluster contains the control plane and worker nodes.",
            "A node is usually a worker machine, often a VM, that runs pods.",
            "A pod is the smallest deployable compute unit in Kubernetes.",
            "A ReplicaSet keeps the requested number of pod replicas running.",
            "A Deployment manages ReplicaSets and enables rolling updates and rollbacks.",
            "A Service gives stable network access to pods.",
            "Common service types are ClusterIP, NodePort and LoadBalancer.",
            "ConfigMaps and Secrets provide configuration values to pods.",
        ],
        "simple_explanation": [
            "Kubernetes exists because running a few containers by hand is manageable, but running hundreds or thousands is not. It automates creation, replacement, scaling and traffic routing.",
            "If you want the quickest possible mental model, remember this chain: Deployment chooses the version, ReplicaSet keeps the correct number of pods alive, Pod is where the container runs, and Service gives the pods a stable address.",
            "The most important hierarchy to remember is this: a Deployment controls ReplicaSets, a ReplicaSet keeps pods alive, and a Service provides stable access to the pods. That single chain solves many exam questions.",
        ],
        "example": [
            "A web application may run three pod replicas. A Service routes traffic to them. When version 2 is deployed, the Deployment creates a new ReplicaSet for the new pods, checks they are healthy and then gradually removes the old ReplicaSet's pods."
        ],
        "common_mistakes": [
            "Confusing a pod with a node.",
            "Saying a Service is the same thing as a microservice business concept.",
            "Forgetting that Deployment sits above ReplicaSet.",
            "Describing rolling updates without mentioning health checks or gradual replacement.",
        ],
        "how_in_exams": [
            "2023-2024 asked for the functionality of Kubernetes components including service types.",
            "2024-2025 main asked how rolling updates work using Deployment and ReplicaSet.",
            "2023-2024 refer/defer asked what rolling upgrades mean and how they work.",
        ],
        "examiner": [
            "Correct hierarchy: Deployment, ReplicaSet, pod, Service.",
            "A clear difference between NodePort, ClusterIP and LoadBalancer.",
            "A step-by-step explanation of zero-downtime updates.",
            "Good use of the words 'replica', 'health', 'traffic' and 'rollback'.",
        ],
        "memory_tricks": [
            "Pod runs, ReplicaSet repeats, Deployment rolls, Service routes.",
            "ClusterIP = inside only.",
            "NodePort = exposed on node port.",
            "LoadBalancer = tied to cloud load balancing.",
        ],
        "quiz": [
            "What is the smallest deployable unit in Kubernetes?",
            "What does a ReplicaSet do?",
            "Why is a Deployment needed on top of a ReplicaSet?",
            "What is the difference between ClusterIP and NodePort?",
            "Why can rolling updates avoid downtime?",
        ],
        "quiz_answers": [
            "A pod.",
            "It keeps the required number of pod replicas running.",
            "Because a Deployment manages change over time, especially rolling updates and rollbacks, rather than just keeping replicas alive.",
            "ClusterIP is for internal cluster access only, while NodePort exposes the service on a port of each node so it can be reached from outside.",
            "Because Kubernetes brings healthy new pods into service before fully removing the old ones.",
        ],
        "exam_question": "With the aid of a simple diagram, explain how Kubernetes enables rolling updates with no downtime using a Deployment object and a ReplicaSet object.",
        "model_answer": [
            "A ReplicaSet ensures that a specified number of identical pod replicas are running. On its own, it is useful for availability because if one pod fails, Kubernetes can create another to maintain the replica count.",
            "A Deployment sits above ReplicaSets and manages change over time. When a new application image is deployed, the Deployment creates a new ReplicaSet for the new version instead of simply deleting all old pods at once. It then brings up new pods gradually and checks they are healthy.",
            "As the new pods become ready, traffic can move across to them and the old ReplicaSet is scaled down step by step. This is why downtime can be avoided. If the new version fails, Kubernetes can roll back by returning to the previous ReplicaSet.",
        ],
        "recap": [
            "For Kubernetes update questions, always think: new ReplicaSet up first, old ReplicaSet down later."
        ],
    },
    {
        "title": "Inter-service Communication: HTTP, Queues, Pub-Sub and Event-Driven Design",
        "what": [
            "This topic covers how distributed services talk to each other using direct synchronous calls, asynchronous queues and event-driven patterns."
        ],
        "why": [
            "Messaging is one of the strongest recurring themes across the papers. If you know these patterns well, you will be prepared for many likely questions."
        ],
        "key_ideas": [
            "Synchronous communication means the caller waits for a reply, as with normal HTTP request-response.",
            "Asynchronous queue-based messaging decouples sender and receiver in time.",
            "Point-to-point queues are usually one message to one consumer.",
            "Publisher-subscriber distributes the message or event to multiple interested consumers.",
            "An event is a fact that something happened, often tied to a state change.",
            "Thick events carry the state data; thin events carry minimal data and often require a follow-up API call.",
            "Eventual consistency is common in event-driven systems because replicas update after the main change happens.",
            "Kafka is a persistent event stream; RabbitMQ is a message broker built strongly around queues and exchanges.",
        ],
        "simple_explanation": [
            "If one service calls another by HTTP, it is like making a phone call and waiting on the line for the answer. That is simple and fast when everything is healthy, but it tightly links the two services.",
            "Use this quick map if the vocabulary feels confusing: HTTP = ask and wait. Queue = hand off work and leave. Pub-sub = announce something to many listeners. Event = a fact that something happened, not a request telling another service what to do.",
            "A queue is more like posting a job into an inbox. The sender drops the message and continues. The consumer processes it when ready. Pub-sub goes a step further: instead of one inbox, the message is distributed to many subscribers who all care about the event.",
        ],
        "example": [
            "An orders service may place a message on a queue for payment processing. When payment succeeds, an event such as PaymentConfirmed may be published so shipping, notifications and reporting services can all react independently."
        ],
        "common_mistakes": [
            "Using 'asynchronous' as if it means only one specific technology.",
            "Confusing a shared queue with pub-sub.",
            "Ignoring the trade-off between resilience and immediate consistency.",
            "Talking about events without making it clear what changed and who reacts to it.",
        ],
        "how_in_exams": [
            "2022-2023 main asked about event-driven sharing without sharing a database and compared shared queue with pub-sub.",
            "2023-2024 asked for a redraw from a message-based architecture into an event-driven one.",
            "2024-2025 main asked for HTTP, message queue and event messaging with an advantage, disadvantage and use-case for each.",
            "2024-2025 resit tested message queue and publisher-subscriber directly.",
            "2023-2024 refer/defer tested synchronous vs asynchronous and thick vs thin events.",
        ],
        "examiner": [
            "A correct choice of pattern for the use-case.",
            "Clear trade-offs, not just benefits.",
            "Good wording around coupling, resilience, latency and consistency.",
            "A clean event or message flow where the steps are easy to follow.",
        ],
        "memory_tricks": [
            "HTTP asks now.",
            "Queue asks later.",
            "Pub-sub tells many.",
            "Thin event carries ID; thick event carries state.",
        ],
        "quiz": [
            "Why is HTTP usually described as synchronous between services?",
            "What advantage does a message queue bring when the consumer is temporarily slow or unavailable?",
            "How is pub-sub different from a single shared queue?",
            "What is a thick event?",
            "Why might an event-driven design lead to eventual consistency?",
        ],
        "quiz_answers": [
            "Because the caller normally waits for the response before it can continue with that interaction.",
            "It can buffer the work so the producer does not have to wait for the consumer to be ready right now.",
            "In pub-sub, multiple subscribers can receive the same event, whereas a single shared queue usually delivers each message to one consumer.",
            "An event that carries useful state data with it, not just a tiny notification.",
            "Because updates are propagated asynchronously, so different services or replicas may be briefly out of sync before they catch up.",
        ],
        "exam_question": "Discuss one advantage, one disadvantage and one suitable use-case for each of the following messaging patterns: HTTP, message queue and event messaging.",
        "model_answer": [
            "HTTP request-response is simple and direct. Its advantage is low conceptual overhead and an immediate reply, so it suits tasks where the caller needs the result straight away, such as requesting current product details. Its disadvantage is tighter coupling because the caller waits for the target service and is affected by its failure or slowness.",
            "A message queue is asynchronous and loosely couples producer and consumer. Its advantage is resilience, because the producer can hand off work and continue even if the consumer is busy. A good use-case is background payment or document processing. Its disadvantage is that the result is not immediate and the overall flow is harder to trace.",
            "Event messaging is useful when one change should notify many services. Its advantage is extensibility: new subscribers can react without changing the producer. A good use-case is an OrderPlaced event that drives shipping, notification and analytics. Its disadvantage is higher architectural complexity and the need to manage eventual consistency carefully.",
        ],
        "recap": [
            "If the answer must choose a pattern, always tie the choice to immediacy, coupling, resilience and the number of consumers."
        ],
    },
    {
        "title": "API Gateway",
        "what": [
            "This topic explains the role of an API gateway as the front door to a distributed system."
        ],
        "why": [
            "API gateway questions are common because the gateway sits at the centre of routing, security and client simplification in many microservice systems."
        ],
        "key_ideas": [
            "An API gateway is a single entry point between clients and backend services.",
            "It can route requests to the correct service based on path or other rules.",
            "It can centralise security concerns such as authentication and rate limiting.",
            "It can perform caching, logging, analytics and response aggregation.",
            "It can reduce client complexity because the client does not need to know every backend address.",
            "It can also become a bottleneck or single point of failure if not designed carefully.",
        ],
        "simple_explanation": [
            "Without a gateway, a client may need to know where every backend service lives and may need to handle many cross-cutting concerns itself. With a gateway, the client talks to one front door and the gateway decides where the traffic should go.",
            "This makes the system easier for clients to consume, but it also means the gateway becomes important infrastructure that must itself be scalable and resilient.",
        ],
        "example": [
            "A shopping web app may call one gateway address. The gateway routes catalogue requests to the product service, checkout requests to the order service and account requests to the identity service. It can also apply the same authentication rules to all of them."
        ],
        "common_mistakes": [
            "Calling the gateway just a router and ignoring security or aggregation.",
            "Listing only advantages with no disadvantages.",
            "Forgetting to tie the benefits to the use-case in the question.",
            "Assuming the gateway removes all architectural complexity.",
        ],
        "how_in_exams": [
            "2022-2023 main asked for a use-case, diagram and explanation of why an API gateway solves problems in microservice architecture.",
            "2023-2024 asked for five properties of an API gateway and how they benefit the design.",
            "2024-2025 main asked for the purpose of using an API gateway in a distributed system.",
        ],
        "examiner": [
            "A clear idea of the problem the gateway solves.",
            "Multiple concrete properties such as routing, auth, rate limiting and aggregation.",
            "At least one risk or downside.",
            "A link back to the chosen use-case.",
        ],
        "memory_tricks": [
            "One front door, many backend rooms.",
            "Gateway handles common rules once, not in every service.",
        ],
        "quiz": [
            "Why can an API gateway simplify client development?",
            "What is one security-related job an API gateway can do?",
            "What does aggregation mean in an API gateway context?",
            "Why can an API gateway become a risk if poorly designed?",
            "Give one use-case where a gateway is especially helpful.",
        ],
        "quiz_answers": [
            "Because the client can talk to one front door instead of knowing the details of many backend services.",
            "It can centralise tasks such as authentication, authorisation or rate limiting.",
            "It means combining results from several backend services into one response for the client.",
            "Because it can become a bottleneck, a single point of failure or a bigger attack surface if it is not designed well.",
            "A shopping or booking application where one front end needs data from many separate backend services.",
        ],
        "exam_question": "With the aid of a relevant use-case, explain why you may recommend an API gateway for a microservice architecture.",
        "model_answer": [
            "I would recommend an API gateway when a client needs to access multiple backend microservices through one controlled front door. For example, in an online retail system, the web application may need product, order and account data. Instead of exposing every service directly, the client can call the gateway.",
            "The gateway can route traffic to the correct service, centralise authentication and rate limiting, and sometimes aggregate multiple backend responses into one client response. This reduces duplication of common concerns and hides backend service locations from the client.",
            "The main caution is that the gateway must itself be resilient and scalable, because otherwise it becomes a bottleneck or single point of failure. A strong exam answer gives both the benefits and that trade-off.",
        ],
        "recap": [
            "Gateway answers should always mention both simplification and centralisation, plus the risk of making the gateway too critical."
        ],
    },
    {
        "title": "HTTPS, TLS, Certificates and Digital Signatures",
        "what": [
            "This topic explains how secure web communication works, including the roles of hashing, encryption, signatures, certificates, the chain of trust and the TLS handshake."
        ],
        "why": [
            "This is the single strongest recurring topic in the past papers. It is also a classic source of lost marks when students mix up the crypto tools."
        ],
        "key_ideas": [
            "HTTPS is HTTP protected by TLS.",
            "TLS aims to provide confidentiality, integrity and authenticity during communication.",
            "Hashing checks integrity; HMAC adds shared-secret authenticity; encryption hides data; digital signatures prove sender authenticity and integrity.",
            "Symmetric encryption uses one shared key and is fast.",
            "Asymmetric cryptography uses a public/private key pair and supports key exchange and signatures.",
            "A digital certificate binds a public key to an identity such as a domain.",
            "The chain of trust lets the browser trust a server certificate through trusted root and intermediate authorities.",
            "The TLS handshake agrees the secure session and creates session keys for the traffic that follows.",
        ],
        "simple_explanation": [
            "A common beginner mistake is to think a certificate itself encrypts the data. It does not. The certificate mainly helps the browser trust that the server's public key really belongs to the server it claims to be.",
            "The easiest way to remember the flow is: certificate = who is this server, handshake = how do we safely agree secrets, session key = what actually locks the conversation.",
            "TLS combines several techniques. The certificate and signature help establish trust, key exchange helps the two sides agree secrets safely, and symmetric session keys are then used for the actual high-speed encryption of the data in transit.",
        ],
        "example": [
            "When a browser visits an online bank over HTTPS, it receives the server's certificate, checks the chain of trust, completes the TLS handshake, creates session keys and then uses those keys to protect the traffic that follows."
        ],
        "common_mistakes": [
            "Saying a digital signature is the same thing as encryption for secrecy.",
            "Saying the certificate itself encrypts the session data.",
            "Forgetting to explain why the browser trusts the certificate chain.",
            "Describing TLS as one single algorithm instead of a process using several techniques.",
        ],
        "how_in_exams": [
            "2022-2023 main asked for certificate creation, deployment and browser trust.",
            "2022-2023 reassessment asked for the HTTPS process and where the crypto techniques are used.",
            "2023-2024 asked about digital certificates and their purpose.",
            "2024-2025 main asked how a certificate protects a web user.",
            "2024-2025 resit asked about a digital signature on a contract.",
            "2023-2024 refer/defer asked about authenticity methods, session keys and the full TLS handshake.",
        ],
        "examiner": [
            "A precise difference between hash, HMAC, encryption, signature and certificate.",
            "A clear ordered process from browser request to secure session.",
            "Correct use of the words 'public key', 'private key', 'certificate authority' and 'session key'.",
            "A short explanation of what each stage achieves, not just a list of buzzwords.",
        ],
        "memory_tricks": [
            "Hash checks.",
            "HMAC checks with a shared secret.",
            "Encryption hides.",
            "Signature proves.",
            "Certificate binds key to identity.",
        ],
        "quiz": [
            "Why is symmetric encryption usually used for the main encrypted session data rather than asymmetric encryption?",
            "What problem does a digital certificate solve?",
            "What is the chain of trust?",
            "What does a digital signature prove?",
            "What happens at a high level during the TLS handshake?",
        ],
        "quiz_answers": [
            "Because symmetric encryption is much faster and is better suited to protecting the main flow of session data.",
            "It helps prove that the public key really belongs to the server or identity it claims to represent.",
            "It is the certificate path from the server certificate through intermediate authorities back to a trusted root authority.",
            "It proves authenticity and integrity, and can also support non-repudiation if the private key is properly controlled.",
            "The client and server exchange hello data, the server sends certificates, trust is checked, shared secret material is agreed and session keys are derived before encrypted traffic starts.",
        ],
        "exam_question": "Explain how HTTPS works from the point a browser requests a web resource to the point a secure session is established.",
        "model_answer": [
            "When the browser connects to the server over HTTPS, the TLS handshake begins. The client and server exchange hello information and the server sends its certificate chain. The browser checks that the certificate is valid for the domain, has not expired and chains back to a trusted root authority through any intermediate certificates.",
            "If the certificate is trusted, the browser can trust the server's public key. The two sides then perform key-agreement steps and create shared secret material from which session keys are derived. The exact maths is not the key exam point; the key point is that these steps allow secure session keys to be created without exposing them directly on the network.",
            "Once the handshake completes, the session keys are used for fast symmetric encryption of the data in transit. This gives confidentiality, while message protection and the earlier certificate validation give integrity and authenticity. A strong answer explains the stages in order and says what each stage achieves.",
        ],
        "recap": [
            "Certificate = trust the server key. TLS handshake = agree the secure session. Session keys = protect the traffic."
        ],
    },
    {
        "title": "Identity and Access Management: Authentication, OAuth, OIDC, JWT and SSO",
        "what": [
            "This topic explains how systems identify users, decide what they are allowed to access and pass identity or access information between components."
        ],
        "why": [
            "IAM questions often look scary because of the terminology, but they become straightforward once you separate authentication, authorisation, tokens and the main login flow."
        ],
        "key_ideas": [
            "Authentication asks 'who are you?'",
            "Authorisation asks 'what are you allowed to do?'",
            "SSO lets the user authenticate once and then access multiple related systems.",
            "OAuth2 is mainly for delegated authorisation.",
            "OpenID Connect adds identity on top of OAuth2.",
            "A JWT usually has header, payload and signature.",
            "An access token and an ID token are not the same thing.",
            "The authorisation code flow is a step-by-step browser and server interaction with an identity provider.",
        ],
        "simple_explanation": [
            "A user logging into a modern web app often does not send a password directly to the app itself. Instead, the app redirects the user to an identity provider such as Auth0, Google or Microsoft. That provider authenticates the user and returns tokens the app can trust.",
            "The shortest memory map is: Authentication = who are you, Authorisation = what can you do, OIDC = the login journey, JWT = the signed note carrying data about the result.",
            "The most important exam distinction is this: OAuth is about delegated access, while OIDC extends that so identity information is available too. JWT is just a token format commonly used in these flows.",
        ],
        "example": [
            "A student portal may redirect the browser to an identity provider, receive an authorisation code, exchange it server-to-server for tokens, create an application session and then display the user's name and profile picture."
        ],
        "common_mistakes": [
            "Saying OAuth by itself is mainly an authentication standard.",
            "Calling JWT encrypted by default.",
            "Mixing up access token and ID token.",
            "Confusing authentication and authorisation in written answers.",
        ],
        "how_in_exams": [
            "2022-2023 reassessment asked for a wide range of IAM terminology including OAuth, OIDC, JWT, SAML and SSO.",
            "2023-2024 asked about JWT structure and the OpenID Connect authorisation flow using a sequence diagram.",
            "Identity themes also appear around cloud and security in older papers.",
        ],
        "examiner": [
            "Accurate terminology.",
            "A correct step sequence for the OIDC authorisation code flow.",
            "A good explanation of what the tokens are for.",
            "Clear distinction between authentication and authorisation.",
        ],
        "memory_tricks": [
            "AuthN = identity; AuthZ = access.",
            "OIDC = OAuth + identity.",
            "JWT = header.payload.signature.",
        ],
        "quiz": [
            "What is the difference between authentication and authorisation?",
            "What does SSO give the user?",
            "Why is OIDC often described as a layer on top of OAuth2?",
            "What are the three parts of a JWT?",
            "What is the job of the identity provider in the authorisation code flow?",
        ],
        "quiz_answers": [
            "Authentication checks who the user is, while authorisation checks what the user is allowed to access or do.",
            "It lets the user sign in once and then access multiple related systems without logging in separately each time.",
            "Because it builds on OAuth2 and adds identity information so authentication is supported as well as delegated access.",
            "Header, payload and signature.",
            "It authenticates the user and issues the code and tokens that the application then uses.",
        ],
        "exam_question": "Using a sequence-diagram style explanation, explain how a user accesses protected resources in a web application using the OpenID Connect authorisation code flow.",
        "model_answer": [
            "The user first tries to access a protected page in the web application. The application checks for an existing session and, if none exists, redirects the browser to the identity provider's authorisation endpoint with values such as client ID, redirect URI and scope.",
            "The identity provider authenticates the user and, if consent is needed, asks the user to approve the requested access. It then redirects the browser back to the application's redirect URI with an authorisation code. The application backend exchanges that code server-to-server for tokens, commonly including an access token and, in OIDC, an ID token.",
            "The application validates the tokens, extracts the identity information it needs, creates its own application session and returns the protected page. Later requests send the session cookie and the application checks whether access is still valid. The key to a high-mark answer is keeping the roles of browser, app server and identity provider clear.",
        ],
        "recap": [
            "If the question says OIDC flow, think redirect, login, code, token exchange, session."
        ],
    },
    {
        "title": "Data Layer: Referential Integrity, Connection Pooling, Read Replicas and Scaling",
        "what": [
            "This topic covers how the data layer stays consistent, supports concurrency and scales when traffic grows."
        ],
        "why": [
            "Database questions often reward simple but precise explanations. Students lose marks when they mix up consistency, scaling and persistence ideas."
        ],
        "key_ideas": [
            "Referential integrity protects relationships between parent and child data using constraints.",
            "Common integrity actions include cascade, restrict and set null.",
            "Connection pooling reuses pre-authenticated database connections.",
            "A pool improves performance and supports concurrent work.",
            "Relational databases are hard to scale horizontally because writes and strong consistency are centralised.",
            "Read replicas scale reads by copying data from a primary to read-only replicas.",
            "Replica lag introduces eventual consistency.",
            "Sharding splits data across nodes; consistent hashing reduces redistribution pain when nodes change.",
        ],
        "simple_explanation": [
            "Referential integrity is about keeping links between related rows valid. If a tutor record points to a course record, the database should stop or control actions that would leave that link broken.",
            "Scaling is a different concern. Even a correct database can struggle under heavy traffic. Connection pooling makes access more efficient, while read replicas help when there are far more reads than writes. Full horizontal scaling of relational systems is harder because you cannot simply let many nodes write freely without consistency problems.",
        ],
        "example": [
            "A retail system may keep one primary relational database for writes and several read replicas for product browsing. Users may briefly see slightly old data on the replicas, but the system can serve many more read requests."
        ],
        "common_mistakes": [
            "Thinking read replicas scale writes as well as reads.",
            "Ignoring replica lag and eventual consistency.",
            "Mixing up primary key and foreign key.",
            "Saying sharding is easy without mentioning data movement and rebalancing challenges.",
        ],
        "how_in_exams": [
            "2023-2024 asked about referential integrity and environment variables.",
            "2024-2025 main asked for database connection pooling benefits and a read-replica scaling pattern.",
            "2022-2023 main asked about relational scaling challenges and then about sharding and consistent hashing.",
        ],
        "examiner": [
            "A practical explanation of the problem first.",
            "A clear statement of what the chosen approach improves and what it does not improve.",
            "Accurate use of terms such as replica, primary, foreign key, connection pool and eventual consistency.",
            "A realistic limitation, not just a benefit list.",
        ],
        "memory_tricks": [
            "Pool reuses connections.",
            "Primary writes, replicas read.",
            "Sharding splits data; hashing chooses the node.",
        ],
        "quiz": [
            "What does referential integrity protect?",
            "Why is connection pooling better than opening a fresh connection for every request?",
            "Why are relational databases hard to scale horizontally?",
            "What does a read replica improve?",
            "What problem does consistent hashing try to reduce?",
        ],
        "quiz_answers": [
            "It protects the validity of relationships between related records, such as parent and child rows.",
            "Because it reuses ready-made connections, reduces setup overhead and supports concurrent work more efficiently.",
            "Because writes, joins and strong consistency are harder to coordinate once the data is spread across multiple nodes.",
            "It improves read capacity and can also improve availability for read traffic.",
            "It reduces the amount of data reassignment needed when nodes are added or removed from a distributed system.",
        ],
        "exam_question": "With the aid of a diagram, explain how read replicas can help a relational database scale and include one benefit and one limitation.",
        "model_answer": [
            "Read replicas are additional read-only copies of the primary relational database. The primary remains the only node that accepts writes, and its changes are copied to the replicas. Application components that mainly read data can send those requests to the replicas instead of overloading the primary.",
            "The main benefit is increased read capacity and improved availability, because read traffic can be spread across several machines. This is useful when a system has far more reads than writes, which is common in web applications.",
            "The main limitation is replica lag. Because updates take time to copy across the network, a replica may briefly return stale data. So this pattern helps with horizontal read scaling, but it does not remove write bottlenecks or strong consistency concerns.",
        ],
        "recap": [
            "Read replicas are a partial scaling pattern, not a full solution to every database scaling problem."
        ],
    },
    {
        "title": "Distributed Transactions, ACID and SAGA",
        "what": [
            "This topic explains why traditional transaction guarantees are straightforward inside one database but difficult across multiple microservices, and how SAGA-style thinking helps."
        ],
        "why": [
            "This is a smaller but still important revision area because it appears in older papers and tests whether you understand the limits of distributed systems."
        ],
        "key_ideas": [
            "ACID stands for Atomicity, Consistency, Isolation and Durability.",
            "ACID works naturally inside one transactional database engine.",
            "In a microservice architecture, one business transaction may span several services and databases.",
            "That makes all-or-nothing guarantees much harder.",
            "A SAGA breaks the large transaction into local transactions plus compensating actions if something later fails.",
            "Distributed designs often accept business-managed risk in exchange for flexibility and availability.",
        ],
        "simple_explanation": [
            "A bank transfer inside one database can be protected by one database transaction. Either both account balances update or neither does. That is straightforward when one engine controls everything.",
            "In a microservice system, the order service, payment service and shipping service may all have separate databases. There is no single simple transaction boundary across them. So designers often use process-based approaches, such as SAGA, instead of pretending the whole system is one database.",
        ],
        "example": [
            "An order may be created, then payment reserved, then shipping started. If shipping fails after payment succeeded, the system may need a compensating action such as cancelling the payment reservation rather than trying to roll back the whole world instantly."
        ],
        "common_mistakes": [
            "Assuming a normal single-database transaction can simply be stretched across all microservices.",
            "Forgetting to describe the failure path as well as the success path.",
            "Ignoring the business risk when eventual consistency is accepted.",
            "Writing ACID as four letters without saying what each one means.",
        ],
        "how_in_exams": [
            "2023-2024 asked for the SAGA orchestration data pattern.",
            "2023-2024 refer/defer asked about ACID and the challenge of transactions in a distributed microservice architecture.",
            "These are historically examinable even though they are not the most visible topics in the current slide pack.",
        ],
        "examiner": [
            "A correct definition of ACID.",
            "A clear explanation of why distributed transactions are harder.",
            "A realistic mitigation such as SAGA with compensation.",
            "Awareness that the solution involves trade-offs and possible business risk.",
        ],
        "memory_tricks": [
            "ACID = all, correct, separate, durable.",
            "SAGA = sequence of local steps with undo where needed.",
        ],
        "quiz": [
            "What does Atomicity mean?",
            "Why is Isolation harder to reason about across multiple services?",
            "What is a compensating action?",
            "Why might a business accept eventual consistency in a distributed process?",
            "What is the core idea of a SAGA?",
        ],
        "quiz_answers": [
            "It means all parts of the transaction succeed together or none of them are treated as complete.",
            "Because there is no single database engine controlling every concurrent operation across all services.",
            "An action that undoes or offsets an earlier local step when a later step fails.",
            "Because the business may prefer the flexibility and scalability of the distributed design if short delays or temporary inconsistency are acceptable.",
            "A SAGA breaks one large business transaction into local steps and uses compensation if something goes wrong later.",
        ],
        "exam_question": "Explain the challenge of implementing transactions in a distributed microservice-based architecture and suggest one approach to deal with it if some business risk is acceptable.",
        "model_answer": [
            "Traditional database transactions assume one engine can enforce all-or-nothing behaviour across the whole operation. In a distributed microservice architecture, one business process may involve several separate services and databases, so there is no single simple transaction boundary.",
            "This creates problems if part of the process succeeds and a later part fails. For example, an order may be recorded and payment reserved before shipping fails. A common approach is to use a SAGA-style process where each service performs a local transaction and, if a later stage fails, a compensating action is triggered to undo or offset the earlier step.",
            "This improves practicality in distributed systems, but it introduces business risk because the system may be only eventually consistent and compensation may not be instantaneous. High-mark answers acknowledge both the approach and the trade-off.",
        ],
        "recap": [
            "When the system is split across services, process management often replaces one big database transaction."
        ],
    },
    {
        "title": "Infrastructure as Code with Terraform",
        "what": [
            "This topic explains how infrastructure can be defined in code using declarative tools such as Terraform rather than being built manually through a console."
        ],
        "why": [
            "It is lower-frequency than the biggest topics, but it is clearly in the current syllabus and appeared in the recent paper."
        ],
        "key_ideas": [
            "Infrastructure as Code means infrastructure is described, created and managed through code.",
            "Declarative code states the desired end result; imperative code states the step-by-step actions.",
            "Terraform uses providers, resources, variables and state.",
            "The state file lets Terraform track what already exists.",
            "Provisioners can run local commands, remote commands or copy files.",
            "A null_resource lets provisioners run without representing a real infrastructure object.",
            "IaC improves repeatability, maintainability and automation.",
        ],
        "simple_explanation": [
            "If you build infrastructure manually in a web portal, the result may work, but it is harder to repeat exactly and harder to rebuild later. Infrastructure as Code turns the design into something versioned, reviewable and reusable.",
            "Terraform is mainly declarative. You describe the machines, networking and other resources you want, and Terraform works out the creation order. That is why it is often favoured over purely imperative scripting for repeatable infrastructure work.",
        ],
        "example": [
            "A Terraform configuration may define a virtual network, subnet, public IP, network interface and virtual machine. A file provisioner may then copy a compose file to the server, and a remote-exec provisioner may install Docker."
        ],
        "common_mistakes": [
            "Talking only about code syntax and not the underlying purpose.",
            "Confusing declarative and imperative styles.",
            "Describing a provisioner without saying when it runs.",
            "Forgetting the role of the state file or null_resource.",
        ],
        "how_in_exams": [
            "2024-2025 main asked about imperative vs declarative programming and why declarative is favoured for IaC.",
            "The same paper also tested null_resource, file provisioner and remote-exec provisioner with line-number references.",
        ],
        "examiner": [
            "A clear reason declarative IaC is attractive.",
            "Accurate definitions of key Terraform building blocks.",
            "A practical explanation of what each provisioner does.",
            "Awareness that line-reference questions want purpose, not copied code.",
        ],
        "memory_tricks": [
            "Imperative says how.",
            "Declarative says what.",
            "State remembers reality.",
            "null_resource runs logic without creating infrastructure.",
        ],
        "quiz": [
            "What is Infrastructure as Code?",
            "Why is declarative IaC often preferred over imperative scripting?",
            "What is the job of Terraform's state file?",
            "What does a file provisioner do?",
            "Why might a null_resource be useful?",
        ],
        "quiz_answers": [
            "It is the practice of describing and managing infrastructure using code rather than manual portal clicks.",
            "Because it is more repeatable and lets the tool handle dependencies and resource ordering for you.",
            "It records what infrastructure exists so Terraform can compare the current state with the desired state in code.",
            "It copies files from the local machine to the target machine.",
            "Because it lets you run provisioners or trigger logic without pretending the action itself is a real infrastructure resource.",
        ],
        "exam_question": "Explain why declarative programming is generally favoured for Infrastructure as Code and briefly explain the purpose of Terraform provisioners and null_resource.",
        "model_answer": [
            "Declarative programming is favoured for Infrastructure as Code because the engineer states the desired target environment rather than writing every creation step in order. This improves repeatability, reduces manual sequencing errors and lets the tool work out dependencies between resources.",
            "In Terraform, provisioners are extra actions attached to resource creation. A file provisioner copies files to a machine, a remote-exec provisioner runs commands on the remote machine, and a local-exec provisioner runs commands on the local machine. These are useful for small deployment tasks but should still be explained in terms of purpose, not just syntax.",
            "A null_resource is a special Terraform resource that creates no real infrastructure but can still host provisioners and trigger logic. It is useful when you need an action, such as copying a file after content changes, without pretending that action is a real cloud resource.",
        ],
        "recap": [
            "If the question gives Terraform code with line numbers, explain the intention of the block, not every character in it."
        ],
    },
]

GLOSSARY_BY_TITLE = {
    "Distributed Systems Foundations": [
        ("Distributed system", "A system that looks like one complete system to the user, even though it is built from several separate networked parts."),
        ("Monolith", "One main application deployed as one unit, often with one main database."),
        ("Microservice", "A smaller service built around one business job or area, such as payments or orders."),
        ("Loose coupling", "Services are not heavily tied to each other's internal details, so one can change without breaking lots of others."),
        ("Functional requirement", "What the system must do."),
        ("Non-functional requirement", "How well the system must do it, for example speed, security or reliability."),
        ("Composite microservice", "A higher-level service that coordinates smaller actions or services to deliver one wider business feature."),
    ],
    "API Basics: JSON, URLs, Endpoints and REST": [
        ("API", "A set of rules that lets one program talk to another."),
        ("URL", "The full web address used to reach a resource."),
        ("Endpoint", "A specific API address and method, such as GET /books/42."),
        ("Path parameter", "A value inside the path that usually identifies the resource itself."),
        ("Query parameter", "A value after the ? used to filter, sort or limit results."),
        ("REST", "A common style of API design built around resources and standard HTTP methods."),
        ("HATEOAS", "A stricter REST idea where the response includes links telling the client what actions are possible next."),
    ],
    "Express Request Handling, Middleware, Response Codes and CORS": [
        ("Route", "An Express rule that maps a method and path to the code that should run."),
        ("Middleware", "Code that runs before or around the final endpoint handler."),
        ("req / res", "The request object and response object used by Express when handling a call."),
        ("next()", "The function a middleware calls to hand control to the next step in the Express chain."),
        ("Status code", "A short HTTP number such as 200 or 404 that tells the client what happened."),
        ("Origin", "The scheme, domain and port taken together."),
        ("CORS", "A browser rule that controls whether one origin is allowed to read resources from another origin."),
    ],
    "API Documentation, OpenAPI and Discoverability": [
        ("API contract", "The agreed shape and rules of the API: paths, methods, parameters, responses and errors."),
        ("OpenAPI", "The standard format commonly used to describe and document web APIs."),
        ("Design-first", "Design the API before writing the implementation."),
        ("API-first", "Write the formal API contract first so teams can build against it."),
        ("Code-first", "Start from the implementation and generate or embed documentation from the code."),
        ("Discoverability", "How easy it is for another developer or team to find and understand how to use the API."),
    ],
    "Cloud Fundamentals: Service Models, Shared Responsibility and Business Value": [
        ("On-premises", "Infrastructure that your organisation owns and runs itself."),
        ("Cloud hosting", "Infrastructure or services rented from a provider instead of owned directly."),
        ("IaaS", "Infrastructure as a Service: the provider gives you core infrastructure, but you still manage a lot yourself."),
        ("PaaS", "Platform as a Service: the provider also manages much of the platform and runtime."),
        ("SaaS", "Software as a Service: you mainly use the finished software rather than manage the platform."),
        ("DBaaS", "Database as a Service: a managed database offering."),
        ("Shared responsibility", "The split of security, maintenance and management duties between customer and provider."),
    ],
    "Resilience, Availability Zones, Load Balancers and Elastic Compute": [
        ("High availability", "Keeping the service running during expected component failures."),
        ("Resilience", "The wider ability to withstand, recover from and adapt to failures."),
        ("Availability zone", "A separate cloud location inside a region used to improve fault tolerance."),
        ("Load balancer", "A component that spreads traffic across multiple backend instances."),
        ("Elastic compute", "The ability to add or remove compute resources as demand changes."),
        ("Horizontal scaling", "Adding more instances."),
        ("Vertical scaling", "Making one existing machine bigger or more powerful."),
    ],
    "Containers, Docker, Compose, Volumes and Environment Variables": [
        ("Virtual machine", "A software version of a full computer with its own operating system."),
        ("Container", "A lightweight isolated runtime for an application that shares the host operating system kernel."),
        ("Docker image", "The packaged blueprint used to create containers."),
        ("Container instance", "A running copy of an image."),
        ("Docker Compose", "A declarative way to define and run a multi-container application."),
        ("Named volume", "Docker-managed persistent storage used by containers."),
        ("Bind mount", "A direct mapping from a host folder to a container folder."),
        ("Environment variable", "A runtime setting passed into an application without hard-coding it in source files."),
    ],
    "Kubernetes Fundamentals and Rolling Updates": [
        ("Cluster", "The whole Kubernetes environment."),
        ("Node", "A machine in the cluster, usually a VM or server, that runs workloads."),
        ("Pod", "The smallest deployable unit in Kubernetes."),
        ("ReplicaSet", "A Kubernetes object that makes sure the required number of pods stay running."),
        ("Deployment", "A Kubernetes object that manages ReplicaSets and supports rolling updates and rollbacks."),
        ("Service", "A stable network entry point for reaching pods."),
        ("ConfigMap", "A place to store non-secret configuration values for pods."),
        ("Secret", "A place to store secret values such as passwords or keys."),
    ],
    "Inter-service Communication: HTTP, Queues, Pub-Sub and Event-Driven Design": [
        ("Synchronous", "The caller waits for the result before continuing."),
        ("Asynchronous", "The sender can continue without waiting for the receiver to finish."),
        ("Queue", "A structure where work or messages wait to be handled, usually by one consumer."),
        ("Publisher-subscriber", "A pattern where one event or message can be delivered to many interested consumers."),
        ("Event", "A record that something happened in the system."),
        ("Thick event", "An event that carries useful state data with it."),
        ("Thin event", "An event that mainly says something happened and often carries only an ID or small amount of data."),
        ("Eventual consistency", "The idea that different parts of the system may be briefly out of sync but become consistent later."),
    ],
    "API Gateway": [
        ("API gateway", "One front door through which clients reach multiple backend services."),
        ("Request routing", "Sending a request to the correct backend service."),
        ("Rate limiting", "Restricting how many requests a client can make in a time window."),
        ("Aggregation", "Combining responses from multiple services into one client response."),
        ("Reverse proxy", "A front component that receives requests on behalf of backend services."),
        ("Single point of failure", "One component whose failure can stop the whole service if it is not made resilient."),
    ],
    "HTTPS, TLS, Certificates and Digital Signatures": [
        ("HTTPS", "HTTP protected by TLS."),
        ("TLS", "The protocol that secures data in transit on the network."),
        ("Symmetric encryption", "Encryption that uses the same shared key to encrypt and decrypt."),
        ("Asymmetric encryption", "Encryption that uses a public key and a private key."),
        ("Digital signature", "A cryptographic proof that helps show who sent the data and whether it changed."),
        ("Certificate", "A signed document that binds a public key to an identity such as a domain name."),
        ("Chain of trust", "The link from the server certificate through intermediate authorities back to a trusted root authority."),
        ("Session key", "A temporary symmetric key used to protect the current secure session."),
    ],
    "Identity and Access Management: Authentication, OAuth, OIDC, JWT and SSO": [
        ("Authentication", "Checking who the user is."),
        ("Authorisation", "Checking what the user is allowed to do."),
        ("SSO", "Single Sign-On: logging in once to access multiple related systems."),
        ("Identity provider", "The service that authenticates the user and issues identity or access data."),
        ("OAuth2", "A standard mainly used for delegated authorisation."),
        ("OpenID Connect", "An identity layer on top of OAuth2 that supports authentication as well."),
        ("JWT", "A token format often used to carry signed identity or access information."),
        ("Access token vs ID token", "An access token is for calling protected resources; an ID token is for identity information about the user."),
    ],
    "Data Layer: Referential Integrity, Connection Pooling, Read Replicas and Scaling": [
        ("Referential integrity", "Rules that keep relationships between related rows valid."),
        ("Foreign key", "A field that points to the primary key of another table."),
        ("Connection pool", "A reusable pool of ready-made database connections."),
        ("Primary database", "The main database instance, usually the one that accepts writes."),
        ("Read replica", "A read-only copy of the primary used to spread read traffic."),
        ("Sharding", "Splitting data across multiple nodes instead of keeping it all on one node."),
        ("Consistent hashing", "A way of assigning data to nodes that reduces disruption when nodes are added or removed."),
    ],
    "Distributed Transactions, ACID and SAGA": [
        ("ACID", "The four classic transaction properties: Atomicity, Consistency, Isolation and Durability."),
        ("Atomicity", "All parts of the transaction happen, or none do."),
        ("Consistency", "The transaction leaves the data in a valid state."),
        ("Isolation", "Transactions do not interfere with each other in unsafe ways."),
        ("Durability", "Once committed, the result remains saved even after failure."),
        ("Distributed transaction", "One business process that spans multiple services or databases."),
        ("Compensating action", "A later action used to undo or offset an earlier step when a distributed process fails."),
        ("SAGA", "A distributed transaction pattern built from local steps plus compensation if something goes wrong."),
    ],
    "Infrastructure as Code with Terraform": [
        ("Infrastructure as Code", "Describing and managing infrastructure through code instead of manual clicks."),
        ("Declarative", "You describe the end state you want."),
        ("Imperative", "You describe the step-by-step actions to take."),
        ("Provider", "The Terraform plugin that talks to the cloud provider API."),
        ("Resource", "A Terraform block representing something to create or manage."),
        ("State file", "Terraform's record of what exists and how it matches the code."),
        ("Provisioner", "An extra Terraform feature that can run commands or copy files during resource work."),
        ("null_resource", "A special Terraform resource that creates no real infrastructure but can still run provisioners and triggers."),
    ],
}

HIGH_VALUE_TOPICS = [
    "TLS, certificates, digital signatures and HTTPS process questions.",
    "Messaging patterns: synchronous vs asynchronous, message queues, pub-sub and event-driven flow.",
    "Monolith vs microservices, including trade-offs and API gateway reasoning.",
    "API design, Express middleware, response handling, CORS and OpenAPI documentation.",
    "Cloud models, resilience, load balancing, availability zones and elastic compute.",
    "Containers, named volumes, environment variables, Kubernetes objects and rolling updates.",
    "Identity and access topics: authentication, OAuth, OpenID Connect, JWT and SSO.",
    "Database topics: referential integrity, connection pooling, read replicas and scaling trade-offs.",
    "Lower-frequency but still worth a second pass: HATEOAS, sharding, consistent hashing, ACID/SAGA and Terraform.",
]

MARK_LOSS = [
    "Answering with a definition only when the question really wants a process or flow.",
    "Ignoring the use-case and writing a generic textbook answer.",
    "Listing advantages but not disadvantages or trade-offs.",
    "Drawing a diagram but not labelling it or referring to it in the written explanation.",
    "Mixing up closely related terms such as authentication vs authorisation, queue vs pub-sub, certificate vs signature, or pod vs deployment.",
    "Missing easy marks from examples, short justifications and response-code explanations.",
    "Writing far too much on low-mark sub-parts and then running out of time on higher-mark sections.",
]

EXAM_TECHNIQUE = {
    "Explain questions": [
        "Start with a one-line definition so the examiner knows you know the term.",
        "Then move into an ordered process: step 1, step 2, step 3.",
        "Say why the mechanism matters or what problem it solves.",
        "Finish with a short example or one limitation if there is time.",
    ],
    "Compare questions": [
        "Use the same comparison headings for both sides, for example cost, scaling, resilience and complexity.",
        "Make the comparison explicit with phrases such as 'whereas', 'in contrast' and 'by comparison'.",
        "Do not write two disconnected mini-essays.",
        "End by saying when each option is more suitable.",
    ],
    "Discuss and evaluate questions": [
        "Give both strengths and weaknesses.",
        "Use criteria such as performance, security, resilience, complexity, cost and maintainability.",
        "Tie every point to the scenario in the question.",
        "End with a judgement. If you do not judge, you are usually describing, not evaluating.",
    ],
    "Diagram questions": [
        "Keep diagrams simple but labelled.",
        "Show direction of flow with arrows.",
        "Number the key stages if you plan to explain them in order.",
        "Reference the diagram in your writing: 'At step 3 the load balancer forwards...'.",
        "Never rely on the diagram alone. Marks usually need both drawing and explanation.",
    ],
    "Using definitions, examples, pros and cons, and conclusions": [
        "Definitions should be short and precise, not a whole paragraph.",
        "Examples should be realistic and linked to the topic, for example an online retailer, booking system or login flow.",
        "Advantages and disadvantages should be paired with context.",
        "Conclusions should answer the question, not just repeat earlier points.",
    ],
    "Writing under time pressure": [
        "Read the command word first.",
        "Use the mark count to judge depth. A rough guide is one developed point per 2-3 marks.",
        "Plan in bullets for 30-60 seconds before writing.",
        "If a question invites a diagram, draw it early and build your answer around it.",
        "Leave perfectionism behind. A complete clear answer usually scores better than an unfinished perfect one.",
    ],
    "Before moving on checklist": [
        "Did I answer the exact command word?",
        "Did I define the core term?",
        "Did I explain the process or mechanism clearly?",
        "Did I use the scenario or use-case in the question?",
        "Did I include at least one trade-off, limitation or condition where relevant?",
        "Did I include a diagram or example if it would help?",
        "Did I give a short closing judgement where needed?",
    ],
}

RAPID_REVIEW_SUMMARY = [
    "Distributed systems are chosen to improve agility, scaling and fault isolation, but they add communication, consistency and operational complexity.",
    "A clean API answer normally needs the correct method, a sensible path/query split and a short justification.",
    "Express middleware is all about ordered request handling before the final route response.",
    "CORS is a browser same-origin issue, not just a generic server security feature.",
    "Cloud questions usually reward clear service-model boundaries and shared-responsibility awareness.",
    "Resilience answers should separate staying up, recovering from failure and scaling under load.",
    "Containers package apps efficiently, but persistence must be handled outside the ephemeral container layer.",
    "Kubernetes questions are easiest when you remember the chain: Deployment -> ReplicaSet -> Pods -> Service.",
    "Security questions need clean distinctions: hash, HMAC, encryption, signature, certificate, handshake.",
    "Identity questions need clean distinctions: authentication, authorisation, OAuth, OIDC, JWT, session.",
]

DEFINITIONS = [
    ["Distributed system", "A system that appears as one system to the user but is built from multiple cooperating networked components."],
    ["Monolith", "One main deployable application containing much of the business logic in a single unit."],
    ["Microservice", "A small independently deployable service built around a business capability."],
    ["Loose coupling", "A design where one component depends as little as possible on the internal details of another."],
    ["Endpoint", "A specific API path and method combination that exposes a resource or action."],
    ["Middleware", "Code that runs between receiving a request and sending a response."],
    ["Origin", "Scheme plus domain plus port."],
    ["CORS", "A browser security mechanism controlling whether one origin may call resources on another origin."],
    ["IaaS", "Infrastructure as a Service: rented infrastructure with high customer control and responsibility."],
    ["PaaS", "Platform as a Service: a managed runtime platform for deploying applications."],
    ["Availability zone", "A separate cloud location within a region used to improve resilience."],
    ["Elastic compute", "The ability to scale computing resources up or down in response to demand."],
    ["Named volume", "Docker-managed persistent storage mounted into a container."],
    ["Pod", "The smallest deployable compute unit in Kubernetes."],
    ["ReplicaSet", "A Kubernetes object that maintains a requested number of pod replicas."],
    ["Deployment", "A Kubernetes object that manages ReplicaSets and supports rolling updates and rollbacks."],
    ["Digital signature", "A cryptographic proof of authenticity and integrity created with a private key and checked with a public key."],
    ["Digital certificate", "A signed document that binds a public key to an identity such as a domain."],
    ["JWT", "A JSON Web Token, commonly structured as header.payload.signature."],
    ["OpenID Connect", "An identity layer on top of OAuth2 that supports authentication as well as delegated authorisation."],
    ["Read replica", "A read-only database copy used to spread read traffic away from the primary database."],
    ["Eventual consistency", "A state where replicas may be temporarily out of date but converge later."],
    ["SAGA", "A distributed transaction approach that uses local transactions plus compensation on failure."],
    ["Terraform", "A declarative Infrastructure as Code tool that uses providers, resources and state."],
]

STARTER_GLOSSARY = [
    ["Distributed system", "One system on the outside, many cooperating parts on the inside."],
    ["Monolith", "One main app deployed as one unit."],
    ["Microservice", "A smaller service focused on one business area."],
    ["API", "The rules that let one program talk to another."],
    ["Endpoint", "A specific API address and method, such as GET /books/42."],
    ["Middleware", "Code that runs before the final endpoint response is sent."],
    ["IaaS vs PaaS", "IaaS means you manage more; PaaS means the provider manages more."],
    ["Container", "A lightweight packaged runtime for an application."],
    ["Pod", "The smallest deployable unit in Kubernetes."],
    ["HTTPS / TLS", "HTTPS is HTTP protected by TLS."],
    ["Authentication vs authorisation", "Who are you? / What are you allowed to do?"],
    ["Queue vs event", "A queue holds work to be processed; an event records that something happened."],
]

MARK_SKELETONS = [
    [
        "5 marks",
        "Usually 2-3 clear points with a short definition and one small example.",
        "1. Define the term. 2. Give two explained points. 3. Add one example or use-case.",
    ],
    [
        "10 marks",
        "A short but developed explanation, often with a process or balanced comparison.",
        "1. Define it. 2. Explain 3-4 developed points or steps. 3. Add one advantage and one limitation or one use-case.",
    ],
    [
        "15 marks",
        "A fuller structured answer, often helped by a diagram and explicit justification.",
        "1. Define the idea. 2. Draw or describe the flow. 3. Explain 4-6 developed points. 4. Link them to the scenario. 5. Finish with a judgement.",
    ],
    [
        "25 marks",
        "A complete exam answer with breadth, depth, structure and usually trade-offs.",
        "1. Spend 30-60 seconds planning. 2. Give a precise definition. 3. Use a labelled diagram if relevant. 4. Explain the process in order. 5. Add pros, cons and use-case links. 6. Finish with a conclusion.",
    ],
]

COMPARE_PAIRS = [
    ["Monolith vs microservices", "Simplicity and tight local performance vs agility, independent scaling and fault isolation."],
    ["HTTP vs HTTPS", "Plain transport vs transport protected by TLS."],
    ["Synchronous vs asynchronous", "Wait for reply now vs hand off work and continue."],
    ["Queue vs pub-sub", "One message to one consumer vs one message/event to many interested consumers."],
    ["VM vs container", "Full OS isolation vs lighter process-level isolation."],
    ["IaaS vs PaaS", "More control and more responsibility vs less control and less day-to-day platform management."],
    ["JWT vs server session cookie", "Self-contained token format vs server-managed session reference."],
    ["NodePort vs ClusterIP vs LoadBalancer", "Externally reachable on node port vs internal only vs integrated with cloud load balancing."],
    ["Named volume vs bind mount", "Docker-managed persistent storage vs direct mapping to a host path."],
    ["Thick vs thin event", "Carries state data vs carries minimal notification data."],
]

DIAGRAMS_TO_REPRODUCE = [
    "Client -> API Gateway -> multiple microservices.",
    "Client -> Load Balancer -> multiple app instances across two availability zones.",
    "Deployment -> old ReplicaSet and new ReplicaSet -> pods during a rolling update.",
    "Browser -> App Server -> Identity Provider -> code -> token exchange -> protected resource.",
    "Browser -> Server certificate/handshake -> session keys -> encrypted traffic.",
    "Primary database -> read replicas -> application read traffic spread across replicas.",
    "Producer -> queue or exchange -> one consumer or many subscribers.",
]

SEQUENCES_TO_MEMORISE = [
    [
        "TLS / HTTPS handshake",
        "ClientHello -> ServerHello plus certificate -> browser validates certificate chain -> key agreement -> session keys derived -> encrypted traffic begins.",
    ],
    [
        "Digital signature verification",
        "Sender hashes message -> signs digest with private key -> receiver decrypts signature with public key -> receiver hashes message -> compare digests.",
    ],
    [
        "OIDC authorisation code flow",
        "User requests protected page -> app redirects to IdP -> user logs in -> IdP returns code -> app exchanges code for tokens -> app creates session -> user accesses protected resource.",
    ],
    [
        "Kubernetes rolling update",
        "Deployment creates new ReplicaSet -> new pods start and pass health checks -> traffic shifts -> old ReplicaSet scales down.",
    ],
    [
        "Event-carried state transfer",
        "Service writes source-of-record change -> publishes event with state -> subscribers consume event -> local copies update -> system reaches eventual consistency.",
    ],
]


def set_font(run, name="Calibri", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 24), ("Subtitle", 14), ("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)


def add_paragraph(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    run.italic = italic
    if align is not None:
        p.alignment = align
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell(hdr[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    return table


def add_topic_section(doc, topic):
    doc.add_page_break()
    doc.add_heading(topic["title"], level=2)

    sections = [
        ("What this topic is", topic["what"], False),
        ("Why it matters", topic["why"], False),
        ("Key ideas I must know for the exam", topic["key_ideas"], True),
        ("Simple explanation from scratch", None, None),
        ("Example", topic["example"], False),
        ("Common mistakes", topic["common_mistakes"], True),
        ("How it appears in past exams", topic["how_in_exams"], True),
        ("What an examiner is looking for", topic["examiner"], True),
        ("Short memory tricks", topic["memory_tricks"], True),
    ]

    for heading, content, bullets in sections:
        doc.add_heading(heading, level=3)
        if heading == "Simple explanation from scratch":
            glossary = GLOSSARY_BY_TITLE.get(topic["title"], [])
            if glossary:
                add_paragraph(
                    doc,
                    "Quick term fixes before we start. These are the words you should understand before you try the quiz:",
                )
                for term, meaning in glossary:
                    p = doc.add_paragraph(style="List Bullet")
                    term_run = p.add_run(f"{term}: ")
                    set_font(term_run, bold=True)
                    meaning_run = p.add_run(meaning)
                    set_font(meaning_run)
            for paragraph in topic["simple_explanation"]:
                add_paragraph(doc, paragraph)
        elif bullets:
            add_bullets(doc, content)
        else:
            for paragraph in content:
                add_paragraph(doc, paragraph)

    doc.add_heading("5 quick quiz questions", level=3)
    add_paragraph(
        doc,
        "These questions are meant to be answerable from the section above. If one feels unfamiliar, go back to the quick term fixes and try again.",
    )
    add_numbered(doc, topic["quiz"])
    add_paragraph(doc, "Quick answers for self-check:", bold=True)
    for idx, answer in enumerate(topic["quiz_answers"], start=1):
        add_paragraph(doc, f"Answer {idx}: {answer}")

    doc.add_heading("1 exam-style question", level=3)
    add_paragraph(doc, topic["exam_question"])

    doc.add_heading("1 model answer", level=3)
    for paragraph in topic["model_answer"]:
        add_paragraph(doc, paragraph)

    doc.add_heading("A short recap", level=3)
    for paragraph in topic["recap"]:
        add_paragraph(doc, paragraph)


def validate_content():
    assert len(REVISION_PLAN) == 21, "Revision plan must cover 21 days."
    assert len(TOPICS) >= 10, "Expected a substantial set of core topics."
    for topic in TOPICS:
        assert len(topic["quiz"]) == 5, f"{topic['title']} must have 5 quiz questions."
        assert len(topic["quiz_answers"]) == 5, f"{topic['title']} must have 5 quiz answers."
        assert topic["exam_question"], f"{topic['title']} missing exam question."
        assert topic["model_answer"], f"{topic['title']} missing model answer."


def build_doc():
    validate_content()
    doc = Document()
    style_document(doc)
    doc.core_properties.title = "Distributed Systems - 3-Week Exam Revision Guide"
    doc.core_properties.subject = "Exam-focused revision guide"
    doc.core_properties.author = "OpenAI Codex"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Distributed Systems")
    set_font(run, size=24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("3-Week Exam Revision Guide")
    set_font(run, size=14, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Built from lecture content and past exam questions")
    set_font(run, italic=True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scope.add_run("Based on the 2025/26 lecture slide deck and six past papers from 2022/23 to 2024/25.")
    set_font(run)

    doc.add_page_break()

    doc.add_heading("How to Use This Document", level=1)
    add_paragraph(doc, "Use this as an active revision workbook, not as something to read once and forget.")
    add_paragraph(
        doc,
        "If you are starting from zero, do not panic if the first read feels slow. Your job on the first pass is to understand the plain-English idea, not to memorise every technical word immediately.",
    )
    add_numbered(
        doc,
        [
            "Read one topic section slowly until the basic idea makes sense.",
            "Close the notes and answer the five quick questions from memory.",
            "Attempt the exam-style question without looking.",
            "Check the model answer and identify what you missed.",
            "Return to the topic later using spaced repetition rather than cramming it once.",
        ],
    )
    add_paragraph(doc, "A good daily routine is: learn -> close notes -> recall -> attempt a question -> review mistakes -> revisit later.")
    doc.add_heading("Starter Glossary", level=2)
    add_paragraph(
        doc,
        "Read these first. They are the terms that appear again and again across the module, and knowing them early makes the rest of the guide much easier to follow.",
    )
    add_table(doc, ["Term", "Plain-English meaning"], STARTER_GLOSSARY)

    doc.add_page_break()

    doc.add_heading("Past Paper Analysis", level=1)
    add_paragraph(
        doc,
        "I analysed six papers: " + ", ".join(PAPERS_ANALYSED) + ". The overall pattern is clear: this module's exam heavily rewards clear explanation, process, justification, use-case mapping and labelled diagrams. It is not mainly a memory-dump exam.",
    )
    add_paragraph(
        doc,
        "One strong pattern stands out: 'explain' is by far the dominant command word across the papers, and diagram-led prompts appear again and again. That means your safest revision strategy is to practise structured written explanations supported by simple labelled diagrams.",
    )
    add_table(
        doc,
        ["Repeated topic group", "How often", "What keeps being tested", "Priority"],
        REPEATED_TOPICS,
    )

    doc.add_heading("Command Words the Examiner Keeps Using", level=2)
    add_paragraph(doc, "The papers are dominated by 'explain' style prompts. The safest exam habit is to define, then walk through the process, then say why it matters.")
    add_table(
        doc,
        ["Command word", "What the examiner wants", "How to structure the answer", "Common mistakes that lose marks"],
        COMMAND_WORDS,
    )

    doc.add_heading("Where Students Are Likely to Drop Marks", level=2)
    add_bullets(doc, MARK_LOSS)

    doc.add_heading("Most Likely / Highest Value Topics", level=2)
    add_numbered(doc, HIGH_VALUE_TOPICS)
    add_paragraph(
        doc,
        "Revision priority note: topics such as HATEOAS, sharding, consistent hashing and ACID/SAGA appeared in older papers but are less central in the current slide deck. They are still worth a second-pass review after the core high-frequency topics above.",
    )

    doc.add_page_break()

    doc.add_heading("3-Week Revision Plan", level=1)
    add_paragraph(
        doc,
        "Aim for two focused study blocks per day plus a short review block. At the end of each day, spend 10-15 minutes revisiting material from 1 day ago, 3 days ago and 7 days ago.",
    )
    add_table(
        doc,
        ["Day", "Topic(s) to study", "Order I should study them", "Realistic session goal", "Quick recall task", "Exam-practice task"],
        [[row["day"], row["topics"], row["order"], row["goal"], row["recall"], row["exam"]] for row in REVISION_PLAN],
    )

    doc.add_page_break()

    doc.add_heading("Core Topic Revision Sections", level=1)
    add_paragraph(doc, "The sections below are written to take you from intuition first to exam-ready phrasing second.")

    for topic in TOPICS:
        add_topic_section(doc, topic)

    doc.add_page_break()

    doc.add_heading("Exam Technique", level=1)
    doc.add_heading("What 5-, 10-, 15- and 25-Mark Answers Look Like", level=2)
    add_table(doc, ["Mark level", "What it usually needs", "Safe structure to use"], MARK_SKELETONS)
    for section_title, items in EXAM_TECHNIQUE.items():
        doc.add_heading(section_title, level=2)
        add_bullets(doc, items)

    doc.add_page_break()

    doc.add_heading("Weak Spot / Spaced Repetition Tracker", level=1)
    add_paragraph(doc, "Fill this in honestly after each revision session or timed question.")
    tracker = doc.add_table(rows=1, cols=5)
    tracker.style = "Table Grid"
    headers = ["Topic", "Confidence /10", "Mistakes made", "When to revisit", "Short improvement action"]
    for idx, header in enumerate(headers):
        set_cell(tracker.rows[0].cells[idx], header, bold=True)
    for _ in range(12):
        row = tracker.add_row().cells
        for cell in row:
            set_cell(cell, "")

    doc.add_page_break()

    doc.add_heading("Final Rapid Review Section", level=1)

    doc.add_heading("Condensed High-Priority Summary of the Module", level=2)
    add_bullets(doc, RAPID_REVIEW_SUMMARY)

    doc.add_heading("Quick-Fire Definitions List", level=2)
    add_table(doc, ["Term", "Quick definition"], DEFINITIONS)

    doc.add_heading("Common Compare / Contrast Pairs", level=2)
    add_table(doc, ["Pair", "What to say quickly"], COMPARE_PAIRS)

    doc.add_heading("Common Diagrams I Should Be Able to Reproduce", level=2)
    add_bullets(doc, DIAGRAMS_TO_REPRODUCE)

    doc.add_heading("Common Security / Message-Flow Sequences I Should Memorise", level=2)
    add_table(doc, ["Sequence", "Core order to remember"], SEQUENCES_TO_MEMORISE)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
